# -*- coding=utf-8 -*-
# @Time: 18/4/2025 下午 3:29
# @Author: 席灏铖
# @File: app.py
# @Software: PyCharm

import os
import re
import shutil
import tempfile
import time
import zipfile
from datetime import datetime, time, timedelta
from io import BytesIO
from sqlite3 import IntegrityError
from urllib.parse import quote

# from flask_migrate import Migrate
import requests
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify
from flask import send_file as sf
from flask_apscheduler import APScheduler
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.filters import AutoFilter
from sqlalchemy import or_
from sqlalchemy import and_
from werkzeug.security import check_password_hash

from config import Config
from email_utils import send_email_with_attachment
from models import db, Project, Bid, init_db, MailLog, Leader, SubProject
from oss import upload_file_to_oss, upload_json_to_oss, delete_file_from_oss
from sync_oss import upload_db_to_oss, sync_static_to_oss
from werkzeug.exceptions import RequestEntityTooLarge
from sqlalchemy.exc import IntegrityError

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 限制最大上传文件为 50MB
app.config.from_object(Config)
db.init_app(app)


# ================== 工具函数 ==================
def is_within_business_hours():
    now = datetime.now().time()
    return time(9, 0) <= now <= time(17, 0)


def now_string():
    return datetime.now().strftime('%Y年%m月%d日%H时%M分')


@app.context_processor
def inject_now_string():
    return dict(now_string=now_string)


@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    flash("❌ 文件过大，最大仅支持 50MB！", "danger")

    # 标记跳转后仍保留 session，不清除
    response = redirect(url_for('admin_panel', preserve_form='1'))
    return response


# todo:支持标段在进行中显示
@app.route('/', methods=['GET', 'POST'])
def index():
    session.pop('bid_form_data', None)
    business_open = is_within_business_hours()
    now = datetime.now()

    # ✅ 获取进行中的项目（无论是否有标段）
    available_projects = Project.query.filter(
        Project.start_time <= now,
        Project.deadline >= now
    ).all()

    if request.method == 'POST' and business_open:
        project_code = request.form.get("project_code", "").strip()
        if not project_code:
            flash("请输入项目编号", "danger")
            return redirect(url_for('index'))

        selected_project = Project.query.filter(Project.code.contains(project_code)).first()
        if not selected_project:
            flash("未找到该项目，请检查编号", "danger")
            return redirect(url_for('index'))

        # 标段选择逻辑
        subproject = None
        if selected_project.is_segmented:
            subproject_id = request.form.get("subproject_id")
            if not subproject_id:
                flash("⚠ 该项目有多个标段，请选择具体标段", "warning")
                return redirect(url_for('index'))

            subproject = SubProject.query.get(subproject_id)
            if not subproject or subproject.project_id != selected_project.id:
                flash("❌ 标段选择无效", "danger")
                return redirect(url_for('index'))

        # 手机验证
        phone = request.form.get("phone", "").strip()
        if not re.match(r'^\d{11}$', phone):
            flash("❌ 手机号码格式不正确，请输入11位有效手机号码", "danger")
            return redirect(url_for('index'))

        # 邮箱验证
        email = request.form.get("email", "").strip()
        if not re.match(r'^[^@]+@[^@]+\.(com)$', email):
            flash("❌ 邮箱格式不正确，必须包含 @ 并以 .com 结尾", "danger")
            return redirect(url_for('index'))

        # 信用代码验证
        credit_code = request.form.get("credit_code", "").strip().upper()
        if not re.match(r'^[0-9A-HJ-NP-RT-UWXY]{18}$', credit_code):
            flash("❌ 统一社会信用代码格式不正确", "danger")
            return redirect(url_for('index'))

        # 整合数据
        bid_data = {
            "project_code": selected_project.code,
            "project_name": selected_project.name,
            "subproject_name": subproject.segment_name if subproject else "无标段",
            "supplier_name": request.form.get("supplier_name"),
            "supplier_address": request.form.get("supplier_address"),
            "legal_person": request.form.get("legal_person"),
            "credit_code": credit_code,
            "agent": request.form.get("agent"),
            "phone": phone,
            "email": email,
            "file_method": request.form.get("file_method"),
            "file_time": now_string()
        }

        session['bid_form_data'] = bid_data

        if 'bid_id' in session:
            bid = Bid.query.get(session['bid_id'])
            if bid:
                bid.supplier_name = bid_data["supplier_name"]
                bid.supplier_address = bid_data["supplier_address"]
                bid.legal_person = bid_data["legal_person"]
                bid.credit_code = bid_data["credit_code"]
                bid.agent = bid_data["agent"]
                bid.phone = bid_data["phone"]
                bid.email = bid_data["email"]
                bid.file_method = bid_data["file_method"]
                bid.file_time = bid_data["file_time"]
                db.session.commit()

                return render_template("success_page.html",
                                       data=bid_data,
                                       message="✅ 信息修改成功，请确认或修改。20秒后自动返回。",
                                       message_type="success")
            else:
                flash("❌ 原投标记录未找到，请重新提交", "danger")
                session.pop('bid_id', None)
                return redirect(url_for('index'))
        else:
            existing_bid = Bid.query.filter_by(
                project_id=selected_project.id,
                supplier_name=bid_data["supplier_name"],
                credit_code=bid_data["credit_code"]
            ).first()

            if existing_bid:
                flash("❌ 该公司已对本项目提交过投标，不能重复提交", "danger")
                return redirect(url_for('index'))

            upload_json_to_oss(bid_data, f"项目JSON/{selected_project.code}_{now_string()}.json")

            new_bid = Bid(
                project_id=selected_project.id,
                sub_project_id=subproject.id if subproject else None,
                supplier_name=bid_data["supplier_name"],
                supplier_address=bid_data["supplier_address"],
                legal_person=bid_data["legal_person"],
                credit_code=bid_data["credit_code"],
                agent=bid_data["agent"],
                phone=bid_data["phone"],
                email=bid_data["email"],
                file_method=bid_data["file_method"],
                file_time=bid_data["file_time"],
                status="pending"
            )
            db.session.add(new_bid)
            db.session.commit()

            session['bid_id'] = new_bid.id

            return render_template("success_page.html",
                                   data=bid_data,
                                   message="✅ 提交成功，请确认或修改。20秒后自动返回。",
                                   message_type="success")

    prefill_data = session.pop('bid_form_data', None)
    return render_template('index.html',
                           projects=available_projects,
                           business_open=business_open,
                           prefill=prefill_data)

# @app.route('/', methods=['GET', 'POST'])
# def index():
#     # 清空 session 数据，防止自动跳转时回显
#     session.pop('bid_form_data', None)
#
#     business_open = is_within_business_hours()
#     now = datetime.now()
#
#     # 获取所有当前处于报名时间内的项目的所有标段
#     active_projects = Project.query.filter(
#         Project.start_time <= now,
#         Project.deadline >= now
#     ).all()
#
#     active_project_ids = [p.id for p in active_projects]
#     # available_projects = Project.query.filter(Project.id.in_(active_project_ids)).all()
#     active_subprojects = SubProject.query.filter(SubProject.project_id.in_(active_project_ids)).all()
#
#     if request.method == 'POST' and business_open:
#         # 获取客户填写的数据
#         project_code = request.form.get("project_code")
#         if not project_code:
#             flash("请输入项目编号", "danger")
#             return redirect(url_for('index'))
#
#         # 执行模糊查询，避免传递空值
#         selected_project = Project.query.filter(Project.code.contains(project_code)).first()
#         if not selected_project:
#             flash("未找到该项目，请检查编号", "danger")
#             return redirect(url_for('index'))
#
#         # 手机号码验证（11位数字）
#         phone = request.form.get("phone")
#         if not re.match(r'^\d{11}$', phone):
#             flash("❌ 手机号码格式不正确，请输入11位有效手机号码", "danger")
#             return redirect(url_for('index'))
#
#         # ✅ 查找该项目的进行中标段
#         active_subprojects = SubProject.query.filter(
#             SubProject.project_id == selected_project.id,
#             SubProject.start_time <= now,
#             SubProject.deadline >= now
#         ).all()
#
#         if not active_subprojects:
#             flash("❌ 当前该项目没有在进行中的标段，无法报名", "danger")
#             return redirect(url_for('index'))
#
#         if len(active_subprojects) > 1:
#             flash("⚠ 该项目存在多个标段，请先选择具体标段再提交", "warning")
#             return redirect(url_for('index'))
#
#         # ✅ 默认选中唯一一个标段
#         subproject = active_subprojects[0]
#
#         # 邮箱验证
#         email = request.form.get("email")
#         if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email) or not email.endswith(".com"):
#             flash("❌ 邮箱格式不正确，必须包含 @ 并以 .com 结尾", "danger")
#             return redirect(url_for('index'))
#
#         # 信用代码格式验证
#         credit_code = request.form.get("credit_code", "").strip().upper()
#         if not re.match(r'^[0-9A-HJ-NP-RT-UWXY]{18}$', credit_code):
#             flash("❌ 统一社会信用代码格式不正确，必须为18位大写字母或数字，且不能包含 I、O、Z、S、V", "danger")
#             return redirect(url_for('index'))
#
#         # 收集客户填写的信息
#         bid_data = {
#             "project_code": selected_project.code,
#             "project_name": selected_project.name,
#             "subproject_name": subproject.segment_name,
#             "supplier_name": request.form.get("supplier_name"),
#             "supplier_address": request.form.get("supplier_address"),
#             "legal_person": request.form.get("legal_person"),
#             "credit_code": credit_code,
#             "agent": request.form.get("agent"),
#             "phone": phone,
#             "email": email,
#             "file_method": request.form.get("file_method"),
#             "file_time": now_string()
#         }
#
#         session['bid_form_data'] = bid_data
#
#         if 'bid_id' in session:
#             # 修改模式
#             bid = Bid.query.get(session['bid_id'])
#             if bid:
#                 bid.supplier_name = bid_data["supplier_name"]
#                 bid.supplier_address = bid_data["supplier_address"]
#                 bid.legal_person = bid_data["legal_person"]
#                 bid.credit_code = bid_data["credit_code"]
#                 bid.agent = bid_data["agent"]
#                 bid.phone = bid_data["phone"]
#                 bid.email = bid_data["email"]
#                 bid.file_method = bid_data["file_method"]
#                 bid.file_time = bid_data["file_time"]
#                 db.session.commit()
#
#                 return render_template("success_page.html",
#                                        data=bid_data,
#                                        message="✅ 信息修改成功，请确认或修改。20秒后自动返回。",
#                                        message_type="success")
#             else:
#                 flash("❌ 原投标记录未找到，请重新提交", "danger")
#                 session.pop('bid_id', None)
#                 return redirect(url_for('index'))
#         else:
#             # 新建模式
#             existing_bid = Bid.query.filter_by(
#                 project_id=selected_project.id,
#                 supplier_name=bid_data["supplier_name"],
#                 credit_code=bid_data["credit_code"]
#             ).first()
#
#             if existing_bid:
#                 flash("❌ 该公司已对本项目提交过投标，不能重复提交！", "danger")
#                 return redirect(url_for('index'))
#
#             upload_json_to_oss(bid_data, f"项目JSON/{selected_project.code}_{now_string()}.json")
#
#             new_bid = Bid(
#                 project_id=selected_project.id,
#                 supplier_name=bid_data["supplier_name"],
#                 supplier_address=bid_data["supplier_address"],
#                 legal_person=bid_data["legal_person"],
#                 credit_code=bid_data["credit_code"],
#                 agent=bid_data["agent"],
#                 phone=bid_data["phone"],
#                 email=bid_data["email"],
#                 file_method=bid_data["file_method"],
#                 file_time=bid_data["file_time"],
#                 status="pending"
#             )
#             db.session.add(new_bid)
#             db.session.commit()
#
#             session['bid_id'] = new_bid.id
#
#             return render_template("success_page.html",
#                                    data=bid_data,
#                                    message="✅ 提交成功，请确认或修改。20秒后自动返回。",
#                                    message_type="success")
#
#     # GET 模式回填
#     prefill_data = session.pop('bid_form_data', None)
#     return render_template('index.html', projects=available_projects, business_open=business_open, prefill=prefill_data)

# @app.route('/', methods=['GET', 'POST'])
# def index():
#     # 清空 session 数据，防止自动跳转时回显
#     session.pop('bid_form_data', None)
#
#     business_open = is_within_business_hours()
#     available_projects = Project.query.filter(Project.deadline > datetime.now()).all()
#
#     if request.method == 'POST' and business_open:
#         # 获取客户填写的数据
#         project_code = request.form.get("project_code")
#         if not project_code:
#             flash("请输入项目编号", "danger")
#             return redirect(url_for('index'))
#
#         # 执行模糊查询，避免传递空值
#         selected_project = Project.query.filter(Project.code.contains(project_code)).first()
#         if not selected_project:
#             flash("未找到该项目，请检查编号", "danger")
#             return redirect(url_for('index'))
#
#         # 手机号码验证（11位数字）
#         phone = request.form.get("phone")
#         if not re.match(r'^\d{11}$', phone):
#             flash("❌ 手机号码格式不正确，请输入11位有效手机号码", "danger")
#             return redirect(url_for('index'))
#         # TODO:新增
#         # 查找“正在进行”的标段
#         now = datetime.now()
#         active_subprojects = SubProject.query.filter(
#             SubProject.project_id == selected_project.id,
#             SubProject.start_time <= now,
#             SubProject.deadline >= now
#         ).all()
#
#         if not active_subprojects:
#             flash("❌ 当前该项目没有在进行中的标段，无法报名", "danger")
#             return redirect(url_for('index'))
#
#         if len(active_subprojects) > 1:
#             flash("⚠ 该项目存在多个标段，请先选择具体标段再提交", "warning")
#             return redirect(url_for('index'))
#
#         # ✅ 默认选中唯一一个标段
#         subproject = active_subprojects[0]
#
#         # 邮箱验证（必须有@，并以 .com 结尾）
#         email = request.form.get("email")
#         if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email) or not email.endswith(".com"):
#             flash("❌ 邮箱格式不正确，必须包含 @ 并以 .com 结尾", "danger")
#             return redirect(url_for('index'))
#
#         # 信用代码格式验证（必须18位，不含 I/O/Z/S/V）
#         credit_code = request.form.get("credit_code", "").strip().upper()
#         if not re.match(r'^[0-9A-HJ-NP-RT-UWXY]{18}$', credit_code):
#             flash("❌ 统一社会信用代码格式不正确，必须为18位大写字母或数字，且不能包含 I、O、Z、S、V", "danger")
#             return redirect(url_for('index'))
#
#         # 收集客户填写的信息
#         bid_data = {
#             "project_code": selected_project.code,
#             "project_name": selected_project.name,
#             "subproject_name": subproject.segment_name,
#             "supplier_name": request.form.get("supplier_name"),
#             "supplier_address": request.form.get("supplier_address"),
#             "legal_person": request.form.get("legal_person"),
#             "credit_code": credit_code,
#             "agent": request.form.get("agent"),
#             "phone": phone,
#             "email": email,
#             "file_method": request.form.get("file_method"),
#             "file_time": now_string()
#         }
#
#         # 保存 session 以便回填
#         session['bid_form_data'] = bid_data
#
#         # 判断是新提交还是修改原来的
#         if 'bid_id' in session:
#             # 🛠 修改模式：根据bid_id更新原记录
#             bid = Bid.query.get(session['bid_id'])
#             if bid:
#                 bid.supplier_name = bid_data["supplier_name"]
#                 bid.supplier_address = bid_data["supplier_address"]
#                 bid.legal_person = bid_data["legal_person"]
#                 bid.credit_code = bid_data["credit_code"]
#                 bid.agent = bid_data["agent"]
#                 bid.phone = bid_data["phone"]
#                 bid.email = bid_data["email"]
#                 bid.file_method = bid_data["file_method"]
#                 bid.file_time = bid_data["file_time"]
#                 db.session.commit()
#
#                 # flash("✅ 信息修改成功，请确认或修改。10秒后自动返回。", "success")
#                 # return render_template("success_page.html", data=bid_data)
#
#                 # ✅ 改为传参方式显示提示
#                 return render_template("success_page.html",
#                                        data=bid_data,
#                                        message="✅ 信息修改成功，请确认或修改。20秒后自动返回。",
#                                        message_type="success")
#
#             else:
#                 flash("❌ 原投标记录未找到，请重新提交", "danger")
#                 session.pop('bid_id', None)
#                 return redirect(url_for('index'))
#
#         else:
#             # 🛠 新建模式：检查是否已提交
#             existing_bid = Bid.query.filter_by(
#                 project_id=selected_project.id,
#                 supplier_name=bid_data["supplier_name"],
#                 credit_code=bid_data["credit_code"]
#             ).first()
#
#             if existing_bid:
#                 flash("❌ 该公司已对本项目提交过投标，不能重复提交！", "danger")
#                 return redirect(url_for('index'))  # 如果重复，返回首页并提示
#
#             # 上传JSON到阿里云网盘
#             upload_json_to_oss(bid_data, f"项目JSON/{selected_project.code}_{now_string()}.json")
#
#             # 插入数据库
#             new_bid = Bid(
#                 project_id=selected_project.id,
#                 supplier_name=bid_data["supplier_name"],
#                 supplier_address=bid_data["supplier_address"],
#                 legal_person=bid_data["legal_person"],
#                 credit_code=bid_data["credit_code"],
#                 agent=bid_data["agent"],
#                 phone=bid_data["phone"],
#                 email=bid_data["email"],
#                 file_method=bid_data["file_method"],
#                 file_time=bid_data["file_time"],
#                 status="pending"
#             )
#             db.session.add(new_bid)
#             db.session.commit()
#
#             session['bid_id'] = new_bid.id  # 保存新提交的bid_id
#         # flash("✅ 提交成功，请确认或修改。10秒后自动返回。", "success")
#         # return render_template("success_page.html", data=bid_data)
#         return render_template("success_page.html",
#                                data=bid_data,
#                                message="✅ 提交成功，请确认或修改。20秒后自动返回。",
#                                message_type="success")
#
#     # GET 时回填
#     prefill_data = session.pop('bid_form_data', None)
#     return render_template('index.html', projects=available_projects, business_open=business_open, prefill=prefill_data)


# ================== 管理员登录 ==================
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if request.form['username'] == Config.ADMIN_USERNAME and \
                check_password_hash(Config.ADMIN_PASSWORD, request.form['password']):
            session['admin'] = True
            return redirect(url_for('admin_panel'))
        flash("用户名或密码错误", "danger")
    return render_template("admin_login.html")


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))


# ================== 管理后台 ==================
# 管理员页面最多显示 60 个项目，并且 超出 60 个项目时不显示,但可以搜索到
@app.route('/admin')
def admin_panel():
    if not session.get('admin'):
        return redirect(url_for('login'))
        # ✅ 用户显式请求清空表单数据
    if request.args.get('clear') == '1':
        session.pop('project_form_data', None)
    # all_projects = Project.query.order_by(Project.deadline.desc()).limit(20).all()
    # return render_template("admin_panel.html", projects=all_projects)
    # 获取排序方式
    sort_mode = request.args.get('sort', 'code_asc')  # 默认按项目编号升序
    page = int(request.args.get('page', 1))  # 当前页
    per_page = 10  # 每页项目数
    max_projects = 60

    query = Project.query

    if sort_mode == 'code_asc':
        query = query.order_by(Project.code.asc())
    elif sort_mode == 'code_desc':
        query = query.order_by(Project.code.desc())
    elif sort_mode == 'created_desc':  # 按添加顺序，默认是id降序
        query = query.order_by(Project.id.desc())

    # 筛选总数（不加 limit）
    full_query = query
    total_count = full_query.count()  # 👈 获取完整筛选后的总数

    # 分页显示最多60条中的每页内容
    query = query.limit(max_projects)
    all_projects = query.all()

    start_index = (page - 1) * per_page
    end_index = start_index + per_page
    projects = all_projects[start_index:end_index]

    return render_template(
        "admin_panel.html",
        projects=projects,
        sort_mode=sort_mode,
        page=page,
        total_count=total_count,  # ✅ 正确的筛选后总数量
        per_page=per_page
    )


# @app.route('/admin/project/<int:project_id>/bids')
# def view_bids(project_id):
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#     project = Project.query.get_or_404(project_id)
#     bids = project.bids
#     return render_template("view_bids.html", project=project, bids=bids)
# todo：含子标段分组
@app.route('/admin/project/<int:project_id>/bids')
def view_bids(project_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    project = Project.query.get_or_404(project_id)
    sub_projects = SubProject.query.filter_by(project_id=project.id).all()

    # 按 sub_project_id 聚合每个标段下的投标信息
    grouped_bids = []
    for sub in sub_projects:
        bids = Bid.query.filter_by(sub_project_id=sub.id).order_by(Bid.file_time.desc()).all()
        grouped_bids.append({
            'subproject': sub,
            'bids': bids
        })

    return render_template("view_bids.html", project=project, grouped_bids=grouped_bids)


# 确认缴费页面
# todo：兼容标段结构
@app.route('/mark_paid/<int:bid_id>', methods=['POST'])
def mark_paid(bid_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    bid = Bid.query.get_or_404(bid_id)

    # 获取关联的标段（子项目）和主项目
    subproject = bid.sub_project
    project = subproject.project if subproject else None

    return render_template("confirm_payment.html", bid=bid, subproject=subproject, project=project)


# @app.route('/mark_paid/<int:bid_id>', methods=['POST'])
# def mark_paid(bid_id):
#     if not session.get('admin'):
#         return redirect(url_for('login'))  # 保留管理员验证
#
#     bid = Bid.query.get_or_404(bid_id)
#     return render_template("confirm_payment.html", bid=bid)


# 管理员修改客户投标记录（GET 显示表单，POST 提交更新）
# todo:支持标段
@app.route('/admin/bid/<int:bid_id>/edit', methods=['GET', 'POST'])
def edit_bid(bid_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    bid = Bid.query.get_or_404(bid_id)
    sub_project_id = bid.sub_project_id
    project_id = bid.sub_project.project_id  # ✅ 从标段反查项目 ID

    if bid.is_paid:
        flash("❌ 该投标记录已缴费，不能修改", "danger")
        return redirect(url_for('view_bids', project_id=project_id))

    existing_bids_raw = Bid.query.filter(Bid.sub_project_id == sub_project_id).all()

    existing_bids = [{
        'id': b.id,
        'supplier_name': b.supplier_name,
        'credit_code': b.credit_code
    } for b in existing_bids_raw]

    if request.method == 'POST':
        supplier_name = request.form['supplier_name'].strip()
        supplier_address = request.form['supplier_address'].strip()
        legal_person = request.form['legal_person'].strip()
        credit_code = request.form['credit_code'].strip().upper()
        agent = request.form['agent'].strip()
        phone = request.form['phone'].strip()
        email = request.form['email'].strip()

        if not re.match(r'^\d{11}$', phone):
            flash("❌ 电话号码格式不正确，应为11位数字", "danger")
            return redirect(url_for('edit_bid', bid_id=bid_id))

        if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email) or not email.endswith(".com"):
            flash("❌ 邮箱格式不正确，必须包含 @ 且以 .com 结尾", "danger")
            return redirect(url_for('edit_bid', bid_id=bid_id))

        if not re.match(r'^[0-9A-HJ-NP-RT-UWXY]{18}$', credit_code):
            flash("❌ 统一社会信用代码格式不正确，必须为18位大写字母或数字，且不能包含 I、O、Z、S、V", "danger")
            return redirect(url_for('edit_bid', bid_id=bid_id))

        same_name = Bid.query.filter(
            Bid.sub_project_id == sub_project_id,
            Bid.supplier_name == supplier_name,
            Bid.id != bid.id
        ).first()

        if same_name:
            flash("❌ 当前标段中已存在该投标单位名称，请勿重复", "danger")
            return redirect(url_for('edit_bid', bid_id=bid_id))

        same_code = Bid.query.filter(
            Bid.sub_project_id == sub_project_id,
            Bid.credit_code == credit_code,
            Bid.id != bid.id
        ).first()

        if same_code:
            flash("❌ 当前标段中已存在该统一社会信用代码，请勿重复", "danger")
            return redirect(url_for('edit_bid', bid_id=bid_id))

        bid.supplier_name = supplier_name
        bid.supplier_address = supplier_address
        bid.legal_person = legal_person
        bid.credit_code = credit_code
        bid.agent = agent
        bid.phone = phone
        bid.email = email
        db.session.commit()

        flash("✅ 投标信息已成功修改", "success")
        return redirect(url_for('view_bids', project_id=project_id))

    return render_template("edit_bid.html", bid=bid, existing_bids=existing_bids)


# @app.route('/admin/bid/<int:bid_id>/edit', methods=['GET', 'POST'])
# def edit_bid(bid_id):
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#     bid = Bid.query.get_or_404(bid_id)
#     project_id = bid.project_id
#
#     # 如果是管理员且该投标记录未缴费，允许修改
#     if bid.is_paid:
#         flash("❌ 该投标记录已缴费，不能修改", "danger")
#         return redirect(url_for('view_bids', project_id=bid.project_id))
#
#     # 获取该项目的所有投标记录
#     existing_bids_raw = Bid.query.filter(Bid.project_id == project_id).all()
#
#     # 只提取必要字段构成列表（便于前端 JSON 使用）
#     existing_bids = [{
#         'id': b.id,
#         'supplier_name': b.supplier_name,
#         'credit_code': b.credit_code
#     } for b in existing_bids_raw]
#
#     if request.method == 'POST':
#         # 获取并清洗输入数据
#         supplier_name = request.form['supplier_name'].strip()
#         supplier_address = request.form['supplier_address'].strip()
#         legal_person = request.form['legal_person'].strip()
#         credit_code = request.form['credit_code'].strip().upper()
#         agent = request.form['agent'].strip()
#         phone = request.form['phone'].strip()
#         email = request.form['email'].strip()
#
#         # ✅ 电话格式校验
#         if not re.match(r'^\d{11}$', phone):
#             flash("❌ 电话号码格式不正确，应为11位数字", "danger")
#             return redirect(url_for('edit_bid', bid_id=bid_id))
#
#         # ✅ 邮箱格式校验
#         if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email) or not email.endswith(".com"):
#             flash("❌ 邮箱格式不正确，必须包含 @ 且以 .com 结尾", "danger")
#             return redirect(url_for('edit_bid', bid_id=bid_id))
#
#         # ✅ 统一社会信用代码格式校验
#         if not re.match(r'^[0-9A-HJ-NP-RT-UWXY]{18}$', credit_code):
#             flash("❌ 统一社会信用代码格式不正确，必须为18位大写字母或数字，且不能包含 I、O、Z、S、V", "danger")
#             return redirect(url_for('edit_bid', bid_id=bid_id))
#
#         # ✅ 重复公司名判断（只要同项目中有相同公司名，排除本身）
#         same_name = Bid.query.filter(
#             Bid.project_id == project_id,
#             Bid.supplier_name == supplier_name,
#             Bid.id != bid.id  # 排除当前正在修改的记录
#         ).first()
#
#         if same_name:
#             flash("❌ 当前项目中已存在该投标单位名称，请勿重复", "danger")
#             return redirect(url_for('edit_bid', bid_id=bid_id))
#
#         # ✅ 重复信用代码判断（只要同项目中有相同信用代码，排除本身）
#         same_code = Bid.query.filter(
#             Bid.project_id == project_id,
#             Bid.credit_code == credit_code,
#             Bid.id != bid.id  # 排除当前正在修改的记录
#         ).first()
#
#         if same_code:
#             flash("❌ 当前项目中已存在该统一社会信用代码，请勿重复", "danger")
#             return redirect(url_for('edit_bid', bid_id=bid_id))
#
#         # ✅ 通过所有验证，更新记录
#         bid.supplier_name = supplier_name
#         bid.supplier_address = request.form['supplier_address']
#         bid.legal_person = request.form['legal_person']
#         bid.credit_code = credit_code
#         bid.agent = request.form['agent']
#         bid.phone = phone
#         bid.email = email
#         db.session.commit()
#
#         flash("✅ 投标信息已成功修改", "success")
#         return redirect(url_for('view_bids', project_id=project_id))
#
#     return render_template("edit_bid.html", bid=bid, existing_bids=existing_bids)


# 管理员删除客户投标记录（POST 请求）
@app.route('/admin/bid/<int:bid_id>/delete', methods=['POST'])
def delete_bid(bid_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    bid = Bid.query.get_or_404(bid_id)
    project_id = bid.project_id
    db.session.delete(bid)
    db.session.commit()
    flash("🗑 投标记录已删除", "success")
    return redirect(url_for('view_bids', project_id=project_id))


# 管理员重新发送招标文件邮件（POST 请求）
@app.route('/admin/bid/<int:bid_id>/resend', methods=['POST'])
def resend_email(bid_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    bid = Bid.query.get_or_404(bid_id)
    project = Project.query.get_or_404(bid.project_id)

    # base_filename = f"{project.code} {project.name} 招标文件"
    base_filename = os.path.splitext(os.path.basename(project.file_path))[0]
    result = send_email_with_attachment(
        recipient_email=bid.email,
        subject=f"【{project.code}】招标文件补发通知",
        body=f"您好，补发《{project.name}》的招标文件，请查收。如有问题请联系本公司。",
        base_filename=base_filename
    )

    if result["status"] == "success":
        flash("📧 邮件已成功重新发送", "success")
    else:
        flash("❌ 邮件发送失败：" + result["error"], "danger")

    return redirect(url_for('view_bids', project_id=project.id))


# @app.route('/add_project', methods=['POST'])
# def add_project():
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#     code = request.form['code'].strip()
#     name = request.form['name'].strip()
#     start_time = datetime.strptime(request.form['start_time'], '%Y-%m-%dT%H:%M')
#     deadline = datetime.strptime(request.form['deadline'], '%Y-%m-%dT%H:%M')
#     if deadline <= start_time + timedelta(minutes=1):
#         flash("❌ 截止时间必须晚于开始时间", "danger")
#         return redirect(url_for('admin_panel'))
#     deposit = float(request.form['deposit_amount'])
#     file = request.files['file_upload']
#     ext = os.path.splitext(file.filename)[1].lower()
#     if ext not in Config.ALLOWED_EXTENSIONS:
#         session['project_form_data'] = request.form.to_dict()
#         flash("❌ 不支持的文件类型，请上传 PDF、Word、ZIP 或 RAR 文件", "danger")
#         return redirect(url_for('admin_panel'))
#     if file.content_length and file.content_length > 50 * 1024 * 1024:
#         session['project_form_data'] = request.form.to_dict()
#         flash("❌ 文件太大，不能超过 50MB", "danger")
#         return redirect(url_for('admin_panel'))
#     leader_name = request.form.get('leader_name', '').strip()
#     leader_email = request.form.get('leader_email', '').strip()
#
#     leader = Leader.query.filter_by(name=leader_name).first()
#     if not leader or leader.email != leader_email:
#         flash("❌ 项目组长姓名与邮箱不匹配，请重新选择", "danger")
#         return redirect(url_for('admin_panel'))
#     member_ids_str = request.form.get('member_ids', '')
#     print("💡 提交的 member_ids_str:", member_ids_str)
#     member_ids = [int(id.strip()) for id in member_ids_str.split(',') if id.strip().isdigit()]
#     print("💡 解析后的 member_ids:", member_ids)
#     members = Leader.query.filter(Leader.id.in_(member_ids)).all()
#     print("💡 获取到的成员:", members)
#     filtered_members = [m for m in members if m.email != leader_email]
#     print("💡 过滤后的成员（不包含组长）:", filtered_members)
#
#     # ✅ 邮箱格式二次校验
#     if not re.match(r'^[^@]+@[^@]+\.com$', leader_email):
#         flash("❌ 项目组长邮箱格式不正确，必须包含@且以.com结尾！", "danger")
#         return redirect(url_for('admin_panel'))
#
#     # ✅ 第一步：在什么都没保存前，先检查编号/名称冲突！
#     conflict_project = Project.query.filter(
#         (Project.code == code) | (Project.name == name)
#     ).first()
#     # print(f"🧪 正在查重：输入编号='{code}', 名称='{name}'")
#
#     if conflict_project:
#         session['project_form_data'] = request.form.to_dict()
#         flash("❌ 已存在相同项目编号或名称，请检查！", "danger")
#         return redirect(url_for('admin_panel'))
#
#     # 创建上传目录，确保文件保存路径正确
#     os.makedirs(os.path.join(Config.BASE_DIR, 'static', 'uploads'), exist_ok=True)
#
#     # 获取文件扩展名并保持文件原后缀
#     # ext = os.path.splitext(file.filename)[1]  # 获取文件扩展名
#     # filename = f"{code} {name} 招标文件{ext}"  # 生成文件名并保持原后缀
#     filename = file.filename
#
#     relative_path = os.path.join("static", "uploads", filename)  # 存储在 static/uploads 中
#     absolute_path = os.path.join(Config.BASE_DIR, relative_path)  # 文件的绝对路径
#
#     # 保存文件到指定路径，
#     file.save(absolute_path)
#
#     # 创建项目并存储文件的相对路径
#     project = Project(
#         name=name,
#         code=code,
#         start_time=start_time,
#         deadline=deadline,
#         deposit_amount=deposit,
#         file_path=relative_path,
#         leader_email=leader_email,
#         members=filtered_members  # ✅ 加上这一行：把多选成员赋值给多对多关系字段
#     )
#
#     try:
#         db.session.add(project)
#         db.session.commit()
#     except IntegrityError:
#         db.session.rollback()
#         flash("❌ 数据库中已存在相同项目编号或名称！", "danger")
#         return redirect(url_for('admin_panel'))
#
#     # ✅ 添加成功后清除 session 数据
#     session.pop('project_form_data', None)
#     # 文件上传阿里云
#     upload_file_to_oss(absolute_path, f"项目附件/{filename}")  # 👈 OSS 上传（覆盖）
#
#     flash("✅ 新项目添加成功", "success")
#     return redirect(url_for('admin_panel'))
from sqlalchemy.exc import IntegrityError
@app.route('/add_project', methods=['POST'])
def add_project():
    if not session.get('admin'):
        return redirect(url_for('login'))

    code = request.form['code'].strip()
    name = request.form['name'].strip()
    is_segmented = request.form.get('is_segmented') == 'true'
    segment_count = int(request.form.get('segment_count', 0))
    purchaser = request.form.get('purchaser', '').strip()
    budget_str = request.form.get('budget_amount', '0').strip()
    deposit_str = request.form.get('deposit_amount', '0').strip()

    # ✅ 时间处理
    try:
        start_time = datetime.strptime(request.form['start_time'], '%Y-%m-%dT%H:%M')
        deadline = datetime.strptime(request.form['deadline'], '%Y-%m-%dT%H:%M')
    except:
        flash("❌ 时间格式错误，请检查", "danger")
        return redirect(url_for('admin_panel'))

    try:
        budget_amount = float(budget_str)
    except ValueError:
        flash("❌ 采购金额格式错误", "danger")
        return redirect(url_for('admin_panel'))

    try:
        deposit = float(deposit_str)
    except ValueError:
        flash("❌ 保证金金额格式错误", "danger")
        return redirect(url_for('admin_panel'))

    leader_email = request.form.get('leader_email', '').strip()
    member_ids_str = request.form.get('member_ids', '')
    member_ids = [int(id.strip()) for id in member_ids_str.split(',') if id.strip().isdigit()]
    members = Leader.query.filter(Leader.id.in_(member_ids)).all()
    filtered_members = [m for m in members if m.email != leader_email]

    # ✅ 主项目文件上传
    file = request.files.get('file_upload')
    relative_path = None
    if file and file.filename:
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in Config.ALLOWED_EXTENSIONS:
            flash("❌ 项目文件类型不支持", "danger")
            return redirect(url_for('admin_panel'))
        filename = file.filename
        relative_path = os.path.join("static", "uploads", filename)
        absolute_path = os.path.join(Config.BASE_DIR, relative_path)
        file.save(absolute_path)
        upload_file_to_oss(absolute_path, f"项目附件/{filename}")

    # ✅ 创建主项目
    project = Project(
        name=name,
        code=code,
        leader_email=leader_email,
        members=filtered_members,
        purchaser=purchaser,
        budget_amount=budget_amount,
        is_segmented=is_segmented,
        start_time=start_time,
        deadline=deadline,
        deposit_amount=deposit,
        file_path=relative_path
    )
    db.session.add(project)
    db.session.flush()

    # ✅ 添加标段
    if is_segmented:
        for i in range(segment_count):
            prefix = f"segment_{i}_"
            try:
                segment_name = request.form.get(prefix + 'name', '').strip() or f"{i + 1}标段"

                deposit_raw = request.form.get(prefix + 'deposit_amount', '0').strip()
                try:
                    sub_deposit = float(deposit_raw)
                except ValueError:
                    flash(f"❌ 第{i + 1}个标段保证金金额格式错误", "danger")
                    return redirect(url_for('admin_panel'))

                sub_file = request.files.get(prefix + 'file_upload')
                if not sub_file or not sub_file.filename:
                    flash(f"❌ 第{i + 1}个标段未上传文件", "danger")
                    return redirect(url_for('admin_panel'))

                ext = os.path.splitext(sub_file.filename)[1].lower()
                if ext not in Config.ALLOWED_EXTENSIONS:
                    flash(f"❌ 第{i + 1}个标段文件类型不支持", "danger")
                    return redirect(url_for('admin_panel'))

                sub_relative_path = os.path.join("static", "uploads", sub_file.filename)
                sub_absolute_path = os.path.join(Config.BASE_DIR, sub_relative_path)
                sub_file.save(sub_absolute_path)
                upload_file_to_oss(sub_absolute_path, f"项目附件/{sub_file.filename}")

                sub = SubProject(
                    project_id=project.id,
                    segment_name=segment_name,
                    deposit_amount=sub_deposit,
                    file_path=sub_relative_path,
                    start_time=start_time,  # 所有标段默认与项目时间一致（或后期修改）
                    deadline=deadline
                )
                db.session.add(sub)

            except Exception as e:
                flash(f"❌ 第{i + 1}个标段添加失败：{e}", "danger")
                return redirect(url_for('admin_panel'))

    try:
        db.session.commit()
        session.pop('project_form_data', None)
        flash("✅ 新项目添加成功", "success")
    except IntegrityError:
        db.session.rollback()
        flash("❌ 数据库错误，可能有编号冲突", "danger")

    return redirect(url_for('admin_panel'))

# @app.route('/add_project', methods=['POST'])
# def add_project():
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#     code = request.form['code'].strip()
#     name = request.form['name'].strip()
#     is_segmented = request.form.get('is_segmented') == 'true'
#     segment_count = int(request.form.get('segment_count', 0))
#     purchaser = request.form.get('purchaser', '').strip()
#     budget_str = request.form.get('budget_amount', '0').strip()
#     deposit_str = request.form.get('deposit_amount', '0').strip()
#
#     # ✅ 时间字段统一处理
#     start_str = request.form.get('start_time', '').strip()
#     deadline_str = request.form.get('deadline', '').strip()
#     try:
#         start_time = datetime.strptime(start_str, '%Y-%m-%dT%H:%M')
#         deadline = datetime.strptime(deadline_str, '%Y-%m-%dT%H:%M')
#     except:
#         flash("❌ 时间格式错误，请检查", "danger")
#         return redirect(url_for('admin_panel'))
#
#     try:
#         budget_amount = float(budget_str)
#         deposit = float(deposit_str)
#     except ValueError:
#         flash("❌ 金额格式错误", "danger")
#         return redirect(url_for('admin_panel'))
#
#     leader_email = request.form.get('leader_email', '').strip()
#     member_ids_str = request.form.get('member_ids', '')
#     member_ids = [int(id.strip()) for id in member_ids_str.split(',') if id.strip().isdigit()]
#     members = Leader.query.filter(Leader.id.in_(member_ids)).all()
#     filtered_members = [m for m in members if m.email != leader_email]
#
#     # ✅ 文件处理（无论是否有标段，主项目都可能有文件）
#     file = request.files.get('file_upload')
#     relative_path = None
#     if file and file.filename:
#         ext = os.path.splitext(file.filename)[1].lower()
#         if ext not in Config.ALLOWED_EXTENSIONS:
#             flash("❌ 项目文件类型不支持", "danger")
#             return redirect(url_for('admin_panel'))
#         filename = file.filename
#         relative_path = os.path.join("static", "uploads", filename)
#         absolute_path = os.path.join(Config.BASE_DIR, relative_path)
#         file.save(absolute_path)
#         upload_file_to_oss(absolute_path, f"项目附件/{filename}")
#
#     # ✅ 创建项目
#     project = Project(
#         name=name,
#         code=code,
#         leader_email=leader_email,
#         members=filtered_members,
#         purchaser=purchaser,
#         budget_amount=budget_amount,
#         is_segmented=is_segmented,
#         start_time=start_time,
#         deadline=deadline,
#         deposit_amount=deposit,
#         file_path=relative_path
#     )
#     db.session.add(project)
#     db.session.flush()
#
#     # ✅ 标段信息（仅在有标段时处理）
#     if is_segmented:
#         for i in range(segment_count):
#             prefix = f"segment_{i}_"
#             try:
#                 segment_name = request.form.get(prefix + 'name', '').strip() or f"{i + 1}标段"
#                 sub_deposit = float(request.form[prefix + 'deposit_amount'])
#                 sub_file = request.files[prefix + 'file_upload']
#                 sub_filename = sub_file.filename
#                 ext = os.path.splitext(sub_filename)[1].lower()
#                 if ext not in Config.ALLOWED_EXTENSIONS:
#                     flash(f"❌ 第{i + 1}个标段文件类型不支持", "danger")
#                     return redirect(url_for('admin_panel'))
#
#                 sub_relative_path = os.path.join("static", "uploads", sub_filename)
#                 sub_absolute_path = os.path.join(Config.BASE_DIR, sub_relative_path)
#                 sub_file.save(sub_absolute_path)
#                 upload_file_to_oss(sub_absolute_path, f"项目附件/{sub_filename}")
#
#                 sub = SubProject(
#                     project_id=project.id,
#                     segment_name=segment_name,
#                     deposit_amount=sub_deposit,
#                     file_path=sub_relative_path
#                 )
#                 db.session.add(sub)
#
#             except Exception as e:
#                 flash(f"❌ 第{i + 1}个标段添加失败：{e}", "danger")
#                 return redirect(url_for('admin_panel'))
#
#     try:
#         db.session.commit()
#         session.pop('project_form_data', None)
#         flash("✅ 新项目添加成功", "success")
#     except IntegrityError:
#         db.session.rollback()
#         flash("❌ 数据库错误，可能有编号冲突", "danger")
#
#     return redirect(url_for('admin_panel'))

# @app.route('/add_project', methods=['POST'])
# def add_project():
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#     # ✅ 获取字段
#     code = request.form['code'].strip()
#     name = request.form['name'].strip()
#     is_segmented = request.form.get('is_segmented') == 'true'
#     segment_count = int(request.form.get('segment_count', 0))
#     purchaser = request.form.get('purchaser', '').strip()
#     budget_str = request.form.get('budget_amount', '0').strip()
#
#     try:
#         budget_amount = float(budget_str) if budget_str else 0.0
#     except ValueError:
#         flash("❌ 采购金额格式不正确", "danger")
#         return redirect(url_for('admin_panel'))
#
#     leader_email = request.form.get('leader_email', '').strip()
#     member_ids_str = request.form.get('member_ids', '')
#     member_ids = [int(id.strip()) for id in member_ids_str.split(',') if id.strip().isdigit()]
#     members = Leader.query.filter(Leader.id.in_(member_ids)).all()
#     filtered_members = [m for m in members if m.email != leader_email]
#
#     if not re.match(r'^[^@]+@[^@]+\.com$', leader_email):
#         flash("❌ 项目组长邮箱格式不正确", "danger")
#         return redirect(url_for('admin_panel'))
#
#     conflict_project = Project.query.filter((Project.code == code) | (Project.name == name)).first()
#     if conflict_project:
#         session['project_form_data'] = request.form.to_dict()
#         flash("❌ 已存在相同项目编号或名称，请检查！", "danger")
#         return redirect(url_for('admin_panel'))
#
#     # ✅ 准备上传目录
#     os.makedirs(os.path.join(Config.BASE_DIR, 'static', 'uploads'), exist_ok=True)
#
#     # ✅ 无标段：先处理时间、文件、金额
#     if not is_segmented:
#         start_str = request.form.get('start_time', '').strip()
#         deadline_str = request.form.get('deadline', '').strip()
#
#         if not start_str or not deadline_str:
#             flash("❌ 项目开始和截止时间不能为空！", "danger")
#             return redirect(url_for('admin_panel'))
#
#         try:
#             start_time = datetime.strptime(start_str, '%Y-%m-%dT%H:%M')
#             deadline = datetime.strptime(deadline_str, '%Y-%m-%dT%H:%M')
#         except ValueError:
#             flash("❌ 时间格式错误，请重新选择", "danger")
#             return redirect(url_for('admin_panel'))
#
#         if deadline <= start_time + timedelta(minutes=1):
#             flash("❌ 截止时间必须晚于开始时间", "danger")
#             return redirect(url_for('admin_panel'))
#
#         deposit = float(request.form['deposit_amount'])
#         file = request.files['file_upload']
#         ext = os.path.splitext(file.filename)[1].lower()
#         if ext not in Config.ALLOWED_EXTENSIONS:
#             flash("❌ 不支持的文件类型", "danger")
#             return redirect(url_for('admin_panel'))
#
#         if file.content_length and file.content_length > 50 * 1024 * 1024:
#             flash("❌ 文件大小不能超过50MB", "danger")
#             return redirect(url_for('admin_panel'))
#
#         filename = file.filename
#         relative_path = os.path.join("static", "uploads", filename)
#         absolute_path = os.path.join(Config.BASE_DIR, relative_path)
#         file.save(absolute_path)
#         upload_file_to_oss(absolute_path, f"项目附件/{filename}")
#
#         # ✅ 创建 Project（带所有字段）
#         project = Project(
#             name=name,
#             code=code,
#             leader_email=leader_email,
#             members=filtered_members,
#             purchaser=purchaser,
#             budget_amount=budget_amount,
#             is_segmented=is_segmented,
#             start_time=start_time,
#             deadline=deadline,
#             deposit_amount=deposit,
#             file_path=relative_path
#         )
#         db.session.add(project)
#
#     else:
#         # ✅ 有标段：先创建 Project（标段填进去）
#         project = Project(
#             name=name,
#             code=code,
#             leader_email=leader_email,
#             members=filtered_members,
#             purchaser=purchaser,
#             budget_amount=budget_amount,
#             is_segmented=is_segmented
#         )
#         db.session.add(project)
#         db.session.flush()  # 只需要获得 project.id
#
#         # ✅ 添加每个标段
#         for i in range(segment_count):
#             prefix = f"segment_{i}_"
#             try:
#                 segment_name = request.form.get(prefix + 'name', '').strip() or f"标段{i+1}"
#
#                 deposit = float(request.form[prefix + 'deposit_amount'])
#                 file = request.files[prefix + 'file_upload']
#                 ext = os.path.splitext(file.filename)[1].lower()
#                 if ext not in Config.ALLOWED_EXTENSIONS:
#                     flash(f"❌ 第{i + 1}个标段文件类型不支持", "danger")
#                     return redirect(url_for('admin_panel'))
#
#                 if file.content_length and file.content_length > 50 * 1024 * 1024:
#                     flash(f"❌ 第{i + 1}个标段文件超过50MB", "danger")
#                     return redirect(url_for('admin_panel'))
#
#                 filename = file.filename
#                 relative_path = os.path.join("static", "uploads", filename)
#                 absolute_path = os.path.join(Config.BASE_DIR, relative_path)
#                 file.save(absolute_path)
#                 upload_file_to_oss(absolute_path, f"项目附件/{filename}")
#
#                 sub = SubProject(
#                     project_id=project.id,
#                     segment_name=segment_name,
#                     deposit_amount=deposit,
#                     file_path=relative_path
#                 )
#                 db.session.add(sub)
#
#             except Exception as e:
#                 flash(f"❌ 第{i + 1}个标段处理失败：{e}", "danger")
#                 return redirect(url_for('admin_panel'))
#
#     # ✅ 尝试提交
#     try:
#         db.session.commit()
#         session.pop('project_form_data', None)
#         flash("✅ 新项目添加成功", "success")
#     except IntegrityError:
#         db.session.rollback()
#         flash("❌ 数据库错误，可能有编号冲突", "danger")
#
#     return redirect(url_for('admin_panel'))


# @app.route('/add_project', methods=['POST'])
# def add_project():
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#     # 获取基本字段
#     code = request.form['code'].strip()
#     name = request.form['name'].strip()
#     is_segmented = request.form.get('is_segmented') == 'true'  # ✅ 是否分标段
#     segment_count = int(request.form.get('segment_count', 0))
#
#     leader_name = request.form.get('leader_name', '').strip()
#     leader_email = request.form.get('leader_email', '').strip()
#     member_ids_str = request.form.get('member_ids', '')
#     member_ids = [int(id.strip()) for id in member_ids_str.split(',') if id.strip().isdigit()]
#     members = Leader.query.filter(Leader.id.in_(member_ids)).all()
#     filtered_members = [m for m in members if m.email != leader_email]
#
#     if not re.match(r'^[^@]+@[^@]+\.com$', leader_email):
#         flash("❌ 项目组长邮箱格式不正确，必须包含@且以.com结尾！", "danger")
#         return redirect(url_for('admin_panel'))
#
#     conflict_project = Project.query.filter((Project.code == code) | (Project.name == name)).first()
#     if conflict_project:
#         session['project_form_data'] = request.form.to_dict()
#         flash("❌ 已存在相同项目编号或名称，请检查！", "danger")
#         return redirect(url_for('admin_panel'))
#     # todo：加上采购人和采购金额
#     # ✅ 读取新增字段
#     purchaser = request.form.get('purchaser', '').strip()
#     budget_str = request.form.get('budget_amount', '0').strip()
#     try:
#         budget_amount = float(budget_str) if budget_str else 0.0
#     except ValueError:
#         flash("❌ 采购金额格式不正确，必须是数字", "danger")
#         return redirect(url_for('admin_panel'))
#
#     # ✅ 创建主项目 Project（不含 start_time 等，只用于聚合）
#     project = Project(
#         name=name,
#         code=code,
#         leader_email=leader_email,
#         members=filtered_members,
#         purchaser=purchaser,
#         budget_amount=budget_amount
#     )
#     # project = Project(
#     #     name=name,
#     #     code=code,
#     #     leader_email=leader_email,
#     #     members=filtered_members
#     # )
#     db.session.add(project)
#     db.session.flush()  # 获取 project.id，用于绑定子标段
#
#     os.makedirs(os.path.join(Config.BASE_DIR, 'static', 'uploads'), exist_ok=True)
#
#     if not is_segmented:
#         # ✅ 默认无标段情况：保留原始逻辑，录入为“默认标段”
#         start_time = datetime.strptime(request.form['start_time'], '%Y-%m-%dT%H:%M')
#         deadline = datetime.strptime(request.form['deadline'], '%Y-%m-%dT%H:%M')
#         if deadline <= start_time + timedelta(minutes=1):
#             flash("❌ 截止时间必须晚于开始时间", "danger")
#             return redirect(url_for('admin_panel'))
#
#         deposit = float(request.form['deposit_amount'])
#         file = request.files['file_upload']
#         ext = os.path.splitext(file.filename)[1].lower()
#         if ext not in Config.ALLOWED_EXTENSIONS:
#             flash("❌ 不支持的文件类型，请上传 PDF、Word、ZIP 或 RAR 文件", "danger")
#             return redirect(url_for('admin_panel'))
#
#         if file.content_length and file.content_length > 50 * 1024 * 1024:
#             flash("❌ 文件太大，不能超过 50MB", "danger")
#             return redirect(url_for('admin_panel'))
#
#         filename = file.filename
#         relative_path = os.path.join("static", "uploads", filename)
#         absolute_path = os.path.join(Config.BASE_DIR, relative_path)
#         file.save(absolute_path)
#         upload_file_to_oss(absolute_path, f"项目附件/{filename}")
#
#         sub = SubProject(
#             project_id=project.id,
#             segment_name="默认标段",
#             start_time=start_time,
#             deadline=deadline,
#             deposit_amount=deposit,
#             file_path=relative_path
#         )
#         db.session.add(sub)
#
#     else:
#         # ✅ 多标段模式：前端提供 segment_0_xx, segment_1_xx... 的字段
#         for i in range(segment_count):
#             prefix = f"segment_{i}_"
#             try:
#                 segment_name = request.form[prefix + 'name'].strip() or f"标段{i + 1}"
#                 start_time = datetime.strptime(request.form[prefix + 'start_time'], '%Y-%m-%dT%H:%M')
#                 deadline = datetime.strptime(request.form[prefix + 'deadline'], '%Y-%m-%dT%H:%M')
#                 if deadline <= start_time + timedelta(minutes=1):
#                     flash(f"❌ 第{i + 1}个标段截止时间早于开始时间！", "danger")
#                     return redirect(url_for('admin_panel'))
#
#                 deposit = float(request.form[prefix + 'deposit_amount'])
#                 file = request.files[prefix + 'file_upload']
#                 ext = os.path.splitext(file.filename)[1].lower()
#                 if ext not in Config.ALLOWED_EXTENSIONS:
#                     flash(f"❌ 第{i + 1}个标段：不支持的文件类型", "danger")
#                     return redirect(url_for('admin_panel'))
#
#                 if file.content_length and file.content_length > 50 * 1024 * 1024:
#                     flash(f"❌ 第{i + 1}个标段文件太大，不能超过 50MB", "danger")
#                     return redirect(url_for('admin_panel'))
#
#                 filename = file.filename
#                 relative_path = os.path.join("static", "uploads", filename)
#                 absolute_path = os.path.join(Config.BASE_DIR, relative_path)
#                 file.save(absolute_path)
#                 upload_file_to_oss(absolute_path, f"项目附件/{filename}")
#
#                 sub = SubProject(
#                     project_id=project.id,
#                     segment_name=segment_name,
#                     start_time=start_time,
#                     deadline=deadline,
#                     deposit_amount=deposit,
#                     file_path=relative_path
#                 )
#                 db.session.add(sub)
#
#             except Exception as e:
#                 flash(f"❌ 第{i + 1}个标段处理失败：{e}", "danger")
#                 return redirect(url_for('admin_panel'))
#
#     try:
#         db.session.commit()
#         session.pop('project_form_data', None)
#         flash("✅ 新项目添加成功", "success")
#     except IntegrityError:
#         db.session.rollback()
#         flash("❌ 数据库错误，可能有编号冲突", "danger")
#
#     return redirect(url_for('admin_panel'))


@app.route('/send_file/<int:bid_id>', methods=['POST'])
def send_file(bid_id):
    if not session.get('admin'):
        return redirect(url_for('login'))
    bid = Bid.query.get_or_404(bid_id)
    bid.is_paid = True
    db.session.commit()

    project = bid.project

    # 生成文件名规则：编号 + 项目名称 + 空格 + 招标文件
    filename = f"{project.code} {project.name} 招标文件.docx"  # 生成新的文件名
    # 使用 static/uploads 中的文件路径
    # attachment_path = os.path.join(Config.BASE_DIR, "static", "uploads", filename)
    # 从文件路径中提取原文件名（用于邮件中显示的文件名）
    base_filename = os.path.splitext(os.path.basename(project.file_path))[0]
    # base_filename = f"{project.code} {project.name} 招标文件"

    result = send_email_with_attachment(
        recipient_email=bid.email,
        subject=f"【{project.code}】招标文件发送通知",
        body=f"您好，附件为《{project.name}》的招标文件，请查收。如有问题请联系本公司。",
        base_filename=base_filename
        # attachment_path=attachment_path  # 传递绝对路径
    )

    if result['status'] == 'success':
        flash("邮件发送成功，客户已标记为已缴费。", "success")
    else:
        flash("❌ 邮件发送失败，请手动补发。", "danger")

    return redirect(url_for('view_bids', project_id=project.id))


@app.route('/edit_project/<int:project_id>', methods=['GET', 'POST'])
def edit_project(project_id):
    project = Project.query.get_or_404(project_id)
    sub_projects = SubProject.query.filter_by(project_id=project.id).all()
    is_segmented = project.is_segmented  # ✅ 是否为多标段项目

    if request.method == 'POST':
        # ✅ 解析主项目信息
        new_name = request.form['name']
        new_code = request.form['code']
        # todo:新属性
        new_purchaser = request.form.get('purchaser', '').strip()
        new_purchase_amount = request.form.get('purchase_amount', '0').strip()

        new_start_time = datetime.strptime(request.form['start_time'], '%Y-%m-%dT%H:%M')
        new_deadline = datetime.strptime(request.form['deadline'], '%Y-%m-%dT%H:%M')
        new_deposit = float(request.form['deposit_amount'])
        new_leader_name = request.form.get("new_leader_name", "").strip()
        new_leader_email = request.form.get("new_leader_email", "").strip()

        if not is_segmented:
            if new_deadline <= new_start_time + timedelta(minutes=1):
                flash("❌ 截止时间必须晚于开始时间", "danger")
                return redirect(url_for('edit_project', project_id=project.id))

        if not re.match(r'^[^@]+@[^@]+\.com$', new_leader_email):
            flash("❌ 项目组长邮箱格式不正确，必须包含@且以.com结尾！", "danger")
            return redirect(url_for('edit_project', project_id=project.id))

        # ✅ 校验采购金额为数字
        try:
            new_purchase_amount = float(new_purchase_amount)
        except ValueError:
            flash("❌ 采购金额格式不正确，必须为数字", "danger")
            return redirect(url_for('edit_project', project_id=project.id))

        # ✅ 成员处理
        member_ids_str = request.form.get('member_ids', '')
        member_ids = [int(id.strip()) for id in member_ids_str.split(',') if id.strip().isdigit()]
        old_member_ids = [m.id for m in project.members]
        members_changed = set(member_ids) != set(old_member_ids)
        group_leader_changed = (
                new_leader_email != project.leader_email or
                new_leader_name.lower() not in project.leader_email.lower()
        )

        # ✅ 重复校验
        conflict_project = Project.query.filter(
            ((Project.code == new_code) | (Project.name == new_name)),
            (Project.id != project.id)
        ).first()
        if conflict_project:
            flash("❌ 修改失败：已有相同项目编号或项目名称！", "danger")
            return redirect(url_for('edit_project', project_id=project.id))

        # ✅ 更新项目字段
        project.name = new_name
        project.code = new_code
        # todo:新属性
        project.purchaser = new_purchaser
        project.purchase_amount = new_purchase_amount

        if not is_segmented:
            project.deposit_amount = new_deposit

        project.leader_email = new_leader_email

        deadline_passed = datetime.now() > new_deadline
        if deadline_passed and (group_leader_changed or members_changed):
            flash("❌ 项目已截止，不允许修改组长或成员，但其他已保存", "warning")
        else:
            members = Leader.query.filter(Leader.id.in_(member_ids)).all()
            filtered_members = [m for m in members if m.email != new_leader_email]
            project.members = filtered_members

        # ✅ 替换项目文件（若有）
        file = request.files.get('file_upload')
        if file and file.filename and not is_segmented:
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in Config.ALLOWED_EXTENSIONS:
                flash("❌ 不支持的文件类型，请上传 PDF、Word、ZIP 或 RAR 文件", "danger")
                return redirect(url_for('edit_project', project_id=project.id))
            if file.content_length and file.content_length > 50 * 1024 * 1024:
                flash("❌ 文件太大，不能超过 50MB", "danger")
                return redirect(url_for('edit_project', project_id=project.id))

            old_filename = os.path.basename(project.file_path) if project.file_path else None
            if old_filename:
                try:
                    os.remove(os.path.join(Config.BASE_DIR, project.file_path))
                except FileNotFoundError:
                    pass
                try:
                    delete_file_from_oss(f"项目附件/{old_filename}")
                except Exception as e:
                    print(f"⚠️ 删除 OSS 文件失败: {e}")

            filename = file.filename
            relative_path = os.path.join("static", "uploads", filename)
            absolute_path = os.path.join(Config.BASE_DIR, relative_path)
            file.save(absolute_path)
            upload_file_to_oss(absolute_path, f"项目附件/{filename}")
            project.file_path = relative_path

        # ✅ 处理标段信息更新（仅在启用标段时进行）
        if is_segmented:
            sub_project_ids = request.form.getlist("sub_project_ids[]")
            updated_ids = set()

            for sub_id in sub_project_ids:
                prefix = f"sub_{sub_id}_"
                seg_name = request.form.get(prefix + "segment_name", "").strip()
                seg_start = request.form.get(prefix + "start_time")
                seg_deadline = request.form.get(prefix + "deadline")
                seg_deposit = request.form.get(prefix + "deposit_amount")
                seg_file = request.files.get(prefix + "file_upload")

                if not seg_name or not seg_start or not seg_deadline or not seg_deposit:
                    continue

                seg_start = datetime.strptime(seg_start, "%Y-%m-%dT%H:%M")
                seg_deadline = datetime.strptime(seg_deadline, "%Y-%m-%dT%H:%M")
                seg_deposit = float(seg_deposit)
                seg_file_path = None

                if seg_file and seg_file.filename:
                    ext = os.path.splitext(seg_file.filename)[1].lower()
                    if ext not in Config.ALLOWED_EXTENSIONS:
                        flash(f"❌ 标段 {seg_name} 文件格式错误", "danger")
                        continue
                    filename = seg_file.filename
                    seg_file_path = os.path.join("static", "uploads", filename)
                    seg_abs_path = os.path.join(Config.BASE_DIR, seg_file_path)
                    seg_file.save(seg_abs_path)
                    upload_file_to_oss(seg_abs_path, f"项目附件/{filename}")

                if sub_id.startswith("new"):
                    new_sub = SubProject(
                        project_id=project.id,
                        segment_name=seg_name,
                        start_time=seg_start,
                        deadline=seg_deadline,
                        deposit_amount=seg_deposit,
                        file_path=seg_file_path
                    )
                    db.session.add(new_sub)
                else:
                    sub = SubProject.query.get(int(sub_id))
                    if sub and sub.project_id == project.id:
                        sub.segment_name = seg_name
                        sub.start_time = seg_start
                        sub.deadline = seg_deadline
                        sub.deposit_amount = seg_deposit
                        if seg_file_path:
                            sub.file_path = seg_file_path
                        updated_ids.add(sub.id)

            # ✅ 删除未保留的标段
            current_ids = {s.id for s in sub_projects}
            to_delete = current_ids - updated_ids
            for sid in to_delete:
                sub = SubProject.query.get(sid)
                if sub:
                    db.session.delete(sub)

            # ✅ 支持新增动态标段（来自 JS 动态添加）
            segment_count = int(request.form.get('segment_count', 0))
            for i in range(segment_count):
                prefix = f"segment_{i}_"
                try:
                    segment_name = request.form[prefix + 'name'].strip() or f"标段{i + 1}"
                    start_time = datetime.strptime(request.form[prefix + 'start_time'], '%Y-%m-%dT%H:%M')
                    deadline = datetime.strptime(request.form[prefix + 'deadline'], '%Y-%m-%dT%H:%M')
                    if deadline <= start_time + timedelta(minutes=1):
                        flash(f"❌ 第{i + 1}个新标段截止时间早于开始时间！", "danger")
                        return redirect(url_for('edit_project', project_id=project.id))

                    deposit = float(request.form[prefix + 'deposit_amount'])
                    file = request.files[prefix + 'file_upload']
                    ext = os.path.splitext(file.filename)[1].lower()
                    if ext not in Config.ALLOWED_EXTENSIONS:
                        flash(f"❌ 第{i + 1}个新标段：不支持的文件类型", "danger")
                        return redirect(url_for('edit_project', project_id=project.id))

                    if file.content_length and file.content_length > 50 * 1024 * 1024:
                        flash(f"❌ 第{i + 1}个新标段文件太大，不能超过 50MB", "danger")
                        return redirect(url_for('edit_project', project_id=project.id))

                    filename = file.filename
                    relative_path = os.path.join("static", "uploads", filename)
                    absolute_path = os.path.join(Config.BASE_DIR, relative_path)
                    file.save(absolute_path)
                    upload_file_to_oss(absolute_path, f"项目附件/{filename}")

                    new_sub = SubProject(
                        project_id=project.id,
                        segment_name=segment_name,
                        start_time=start_time,
                        deadline=deadline,
                        deposit_amount=deposit,
                        file_path=relative_path
                    )
                    db.session.add(new_sub)

                except Exception as e:
                    flash(f"❌ 第{i + 1}个新标段处理失败：{e}", "danger")
                    return redirect(url_for('edit_project', project_id=project.id))

        # ✅ 提交更新
        try:
            db.session.commit()
        except IntegrityError:
            db.session.rollback()
            flash("❌ 修改提交失败：数据库已存在冲突！", "danger")
            return redirect(url_for('edit_project', project_id=project.id))

        flash("✅ 项目信息与标段更新成功", "success")
        return redirect(url_for('edit_project', project_id=project.id))

    # ✅ GET 请求渲染页面
    leader = Leader.query.filter_by(email=project.leader_email).first()
    leader_name = leader.name if leader else ""
    member_list = [{'id': m.id, 'name': m.name, 'email': m.email} for m in project.members]
    sub_projects = SubProject.query.filter_by(project_id=project.id).all()

    return render_template("edit_project.html",
                           project=project,
                           is_segmented=is_segmented,
                           now=datetime.now(),
                           leader_name=leader_name,
                           member_list=member_list,
                           sub_projects=sub_projects)


# @app.route('/edit_project/<int:project_id>', methods=['GET', 'POST'])
# def edit_project(project_id):
#     project = Project.query.get_or_404(project_id)
#
#     if request.method == 'POST':
#         new_name = request.form['name']
#         new_code = request.form['code']
#         new_start_time = datetime.strptime(request.form['start_time'], '%Y-%m-%dT%H:%M')
#         new_deadline = datetime.strptime(request.form['deadline'], '%Y-%m-%dT%H:%M')
#
#         # ✅ 开始与截止时间合法性校验
#         if new_deadline <= new_start_time + timedelta(minutes=1):
#             flash("❌ 截止时间必须晚于开始时间", "danger")
#             return redirect(url_for('edit_project', project_id=project.id))
#
#         new_deposit = float(request.form['deposit_amount'])
#         new_leader_name = request.form.get("new_leader_name", "").strip()
#         new_leader_email = request.form.get("new_leader_email", "").strip()
#
#         # ✅ 邮箱格式校验
#         if not re.match(r'^[^@]+@[^@]+\.com$', new_leader_email):
#             flash("❌ 项目组长邮箱格式不正确，必须包含@且以.com结尾！", "danger")
#             return redirect(url_for('edit_project', project_id=project.id))
#
#         # ✅ 解析成员
#         member_ids_str = request.form.get('member_ids', '')
#         member_ids = [int(id.strip()) for id in member_ids_str.split(',') if id.strip().isdigit()]
#         old_member_ids = [m.id for m in project.members]
#         members_changed = set(member_ids) != set(old_member_ids)
#
#         # ✅ 解析组长变更
#         group_leader_changed = (
#                 new_leader_email != project.leader_email or
#                 new_leader_name.lower() not in project.leader_email.lower()
#         )
#
#         # ✅ 检查编号/名称冲突
#         conflict_project = Project.query.filter(
#             ((Project.code == new_code) | (Project.name == new_name)),
#             (Project.id != project.id)
#         ).first()
#         if conflict_project:
#             flash("❌ 修改失败：已有相同项目编号或项目名称！", "danger")
#             return redirect(url_for('edit_project', project_id=project.id))
#
#         # ✅ 允许任何修改（包括时间）
#         project.name = new_name
#         project.code = new_code
#         project.start_time = new_start_time
#         project.deadline = new_deadline
#         project.deposit_amount = new_deposit
#         project.leader_email = new_leader_email
#
#         # ✅ 根据新截止时间判断是否截止
#         deadline_passed = datetime.now() > new_deadline
#
#         # ✅ 截止后拦截组长/成员变更，但允许其他保存
#         if deadline_passed and (group_leader_changed or members_changed):
#             flash("❌ 项目已截止，不允许修改组长或成员，但其他已保存", "warning")
#         else:
#             # ✅ 允许修改成员
#             members = Leader.query.filter(Leader.id.in_(member_ids)).all()
#             filtered_members = [m for m in members if m.email != new_leader_email]
#             project.members = filtered_members
#
#         # ✅ 处理文件上传
#         file = request.files.get('file_upload')
#         if file:
#             ext = os.path.splitext(file.filename)[1].lower()
#             if ext not in Config.ALLOWED_EXTENSIONS:
#                 flash("❌ 不支持的文件类型，请上传 PDF、Word、ZIP 或 RAR 文件", "danger")
#                 return redirect(url_for('edit_project', project_id=project.id))
#             if file.content_length and file.content_length > 50 * 1024 * 1024:
#                 flash("❌ 文件太大，不能超过 50MB", "danger")
#                 return redirect(url_for('edit_project', project_id=project.id))
#
#             old_filename = os.path.basename(project.file_path) if project.file_path else None
#             if project.file_path:
#                 try:
#                     os.remove(os.path.join(Config.BASE_DIR, project.file_path))
#                 except FileNotFoundError:
#                     pass
#
#             filename = file.filename
#             relative_path = os.path.join("static", "uploads", filename)
#             absolute_path = os.path.join(Config.BASE_DIR, relative_path)
#             file.save(absolute_path)
#
#             if old_filename:
#                 try:
#                     delete_file_from_oss(f"项目附件/{old_filename}")
#                 except Exception as e:
#                     print(f"⚠️ 删除 OSS 文件失败: {e}")
#
#             upload_file_to_oss(absolute_path, f"项目附件/{filename}")
#             project.file_path = relative_path
#
#         # ✅ 提交更改
#         try:
#             db.session.commit()
#         except IntegrityError:
#             db.session.rollback()
#             flash("❌ 修改提交失败：数据库已存在冲突！", "danger")
#             return redirect(url_for('edit_project', project_id=project.id))
#
#         flash("✅ 项目更新成功！", "success")
#         return redirect(url_for('edit_project', project_id=project.id))
#
#     # ✅ GET 加载页面
#     leader = Leader.query.filter_by(email=project.leader_email).first()
#     leader_name = leader.name if leader else ""
#     member_list = [{'id': m.id, 'name': m.name, 'email': m.email} for m in project.members]
#     return render_template("edit_project.html", project=project, now=datetime.now(), leader_name=leader_name,
#                            member_list=member_list)


# @app.route('/delete_project/<int:project_id>', methods=['POST'])
# def delete_project(project_id):
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#         # ===== 🟡 第一步：导出当前邮件日志并保存备份 =====
#         logs = MailLog.query.order_by(MailLog.sent_at.desc()).all()
#
#         data = []
#         for idx, log in enumerate(logs, 1):
#             data.append({
#                 '序号': idx,
#                 '发送时间': log.sent_at.strftime('%Y-%m-%d %H:%M:%S') if log.sent_at else '',
#                 '项目名称': log.project.name if log.project else '无关联项目',
#                 '项目编号': log.project.code if log.project else '无编号',
#                 '状态': '成功' if log.status == 'success' else '失败',
#                 '信息': log.message or '无',
#             })
#
#         df = pd.DataFrame(data)
#
#         timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
#         filename = f'邮件发送记录_自动备份_{timestamp}.xlsx'
#         save_dir = os.path.join('static', 'downloads', 'Mail_logs')
#         os.makedirs(save_dir, exist_ok=True)
#         save_path = os.path.join(save_dir, filename)
#         df.to_excel(save_path, index=False, sheet_name='MailLogs', engine='openpyxl')
#
#         # ===== 🔴 第二步：先删除 mail_logs，再删除项目 =====
#         # 删除项目前，先删邮件日志
#     MailLog.query.filter_by(project_id=project_id).delete()
#     project = Project.query.get_or_404(project_id)
#     db.session.delete(project)
#     db.session.commit()
#     flash("项目已删除！", "success")
#     return redirect(url_for('admin_panel'))
# todo：删除项目
@app.route('/delete_project/<int:project_id>', methods=['POST'])
def delete_project(project_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    project = Project.query.get_or_404(project_id)

    try:
        # ===== 1. 当前项目邮件日志导出 =====
        logs = MailLog.query.filter_by(project_id=project_id).order_by(MailLog.sent_at.desc()).all()
        data = []
        for idx, log in enumerate(logs, 1):
            data.append({
                '序号': idx,
                '发送时间': log.sent_at.strftime('%Y-%m-%d %H:%M:%S') if log.sent_at else '',
                '项目名称': log.project.name if log.project else '无关联项目',
                '项目编号': log.project.code if log.project else '无编号',
                '状态': '成功' if log.status == 'success' else '失败',
                '信息': log.message or '无',
            })

        df = pd.DataFrame(data)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_project_name = ''.join(c for c in project.name if c.isalnum() or c in "（）() _-")
        save_dir = os.path.join('static', 'downloads', 'Mail_logs')
        os.makedirs(save_dir, exist_ok=True)

        filename = f'{safe_project_name}_邮件发送记录_自动备份_{timestamp}.xlsx'
        save_path = os.path.join(save_dir, filename)
        df.to_excel(save_path, index=False, sheet_name='MailLogs', engine='openpyxl')

        # ===== 2. 全部邮件日志导出 =====
        all_logs = MailLog.query.order_by(MailLog.sent_at.desc()).all()
        all_data = []
        for idx, log in enumerate(all_logs, 1):
            all_data.append({
                '序号': idx,
                '发送时间': log.sent_at.strftime('%Y-%m-%d %H:%M:%S') if log.sent_at else '',
                '项目名称': log.project.name if log.project else '无关联项目',
                '项目编号': log.project.code if log.project else '无编号',
                '状态': '成功' if log.status == 'success' else '失败',
                '信息': log.message or '无',
            })

        all_df = pd.DataFrame(all_data)
        all_filename = f'全局_邮件记录_完整备份_{timestamp}.xlsx'
        all_path = os.path.join(save_dir, all_filename)
        all_df.to_excel(all_path, index=False, sheet_name='AllMailLogs', engine='openpyxl')

    except Exception as e:
        flash(f"⚠️ 邮件记录导出失败（但项目仍将删除）：{str(e)}", "warning")
        filename = "【导出失败】"
        all_filename = "【全局备份失败】"

    # ===== 3. 删除邮件日志记录 =====
    MailLog.query.filter_by(project_id=project_id).delete()

    # ===== 4. 删除子标段及其文件 =====
    sub_projects = SubProject.query.filter_by(project_id=project_id).all()
    for sub in sub_projects:
        if sub.file_path:
            try:
                os.remove(os.path.join(Config.BASE_DIR, sub.file_path))
            except FileNotFoundError:
                pass
            try:
                delete_file_from_oss(f"项目附件/{os.path.basename(sub.file_path)}")
            except Exception as e:
                print(f"⚠️ 删除 OSS 子标段文件失败: {e}")
        db.session.delete(sub)

    # ===== 5. 删除主项目文件（如有） =====
    if project.file_path:
        try:
            os.remove(os.path.join(Config.BASE_DIR, project.file_path))
        except FileNotFoundError:
            pass
        try:
            delete_file_from_oss(f"项目附件/{os.path.basename(project.file_path)}")
        except Exception as e:
            print(f"⚠️ 删除 OSS 主项目文件失败: {e}")

    # ===== 6. 删除主项目记录 =====
    db.session.delete(project)
    db.session.commit()

    flash(f"✅ 项目删除成功\n📁 单项目备份：{filename}\n📁 全局备份：{all_filename}", "success")
    return redirect(url_for('admin_panel'))


# @app.route('/delete_project/<int:project_id>', methods=['POST'])
# def delete_project(project_id):
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#
#     project = Project.query.get_or_404(project_id)
#
#     try:
#         # ===== 1. 当前项目邮件日志导出 =====
#         logs = MailLog.query.filter_by(project_id=project_id).order_by(MailLog.sent_at.desc()).all()
#         data = []
#         for idx, log in enumerate(logs, 1):
#             data.append({
#                 '序号': idx,
#                 '发送时间': log.sent_at.strftime('%Y-%m-%d %H:%M:%S') if log.sent_at else '',
#                 '项目名称': log.project.name if log.project else '无关联项目',
#                 '项目编号': log.project.code if log.project else '无编号',
#                 '状态': '成功' if log.status == 'success' else '失败',
#                 '信息': log.message or '无',
#             })
#
#         df = pd.DataFrame(data)
#         timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
#         safe_project_name = ''.join(c for c in project.name if c.isalnum() or c in "（）() _-")
#         save_dir = os.path.join('static', 'downloads', 'Mail_logs')
#         os.makedirs(save_dir, exist_ok=True)
#
#         filename = f'{safe_project_name}_邮件发送记录_自动备份_{timestamp}.xlsx'
#         save_path = os.path.join(save_dir, filename)
#         df.to_excel(save_path, index=False, sheet_name='MailLogs', engine='openpyxl')
#
#         # ===== 2. 全部邮件日志导出 =====
#         all_logs = MailLog.query.order_by(MailLog.sent_at.desc()).all()
#         all_data = []
#         for idx, log in enumerate(all_logs, 1):
#             all_data.append({
#                 '序号': idx,
#                 '发送时间': log.sent_at.strftime('%Y-%m-%d %H:%M:%S') if log.sent_at else '',
#                 '项目名称': log.project.name if log.project else '无关联项目',
#                 '项目编号': log.project.code if log.project else '无编号',
#                 '状态': '成功' if log.status == 'success' else '失败',
#                 '信息': log.message or '无',
#             })
#
#         all_df = pd.DataFrame(all_data)
#         all_filename = f'全局_邮件记录_完整备份_{timestamp}.xlsx'
#         all_path = os.path.join(save_dir, all_filename)
#         all_df.to_excel(all_path, index=False, sheet_name='AllMailLogs', engine='openpyxl')
#
#     except Exception as e:
#         flash(f"⚠️ 邮件记录导出失败（但项目仍将删除）：{str(e)}", "warning")
#         filename = "【导出失败】"
#         all_filename = "【全局备份失败】"
#
#     # 删除该项目相关日志
#     MailLog.query.filter_by(project_id=project_id).delete()
#
#     # 删除项目
#     db.session.delete(project)
#     db.session.commit()
#
#     flash(f"✅ 项目删除成功\n📁 单项目备份：{filename}\n📁 全局备份：{all_filename}", "success")
#     return redirect(url_for('admin_panel'))
# TODO: 单独删除某个标段（不会影响主项目或其他标段
@app.route('/delete_sub_project/<int:sub_project_id>', methods=['POST'])
def delete_sub_project(sub_project_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    subproject = SubProject.query.get_or_404(sub_project_id)
    project_id = subproject.project_id  # 保留跳转用

    # 删除标段文件（本地和 OSS）
    if subproject.file_path:
        try:
            os.remove(os.path.join(Config.BASE_DIR, subproject.file_path))
        except FileNotFoundError:
            pass
        try:
            delete_file_from_oss(f"项目附件/{os.path.basename(subproject.file_path)}")
        except Exception as e:
            print(f"⚠️ 删除 OSS 标段文件失败: {e}")

    # 删除标段记录
    db.session.delete(subproject)
    try:
        db.session.commit()
        flash("✅ 标段删除成功", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"❌ 标段删除失败：{e}", "danger")

    return redirect(url_for('edit_project', project_id=project_id))


# 搜索项目，区分管理员和客户
'''
区分管理员和客户的权限：

客户只能查看当前有效的（未过截止时间的）项目。

管理员可以查看所有项目，包括已经过期的项目。
'''


# todo:支持多标段
@app.route('/search_projects')
def search_projects():
    query = request.args.get("query", "").strip()

    if session.get('admin'):  # ✅ 管理员视角
        if query:
            projects = Project.query.filter(
                or_(
                    Project.code.contains(query),
                    Project.name.contains(query)
                )
            ).all()
        else:
            projects = Project.query.all()

        results = []
        for p in projects:
            subprojects = SubProject.query.filter_by(project_id=p.id).all()
            sub_list = []
            for sp in subprojects:
                sub_list.append({
                    'id': sp.id,
                    'name': sp.segment_name or "默认标段",
                    # 这些时间字段建议你从项目表中取
                    'deposit_amount': sp.deposit_amount,
                    'file_path': sp.file_path,
                    'view_bids_url': url_for('view_bids', project_id=p.id, sub_id=sp.id),
                    'edit_project_url': url_for('edit_sub_project', sub_id=sp.id),
                    'delete_project_url': url_for('delete_sub_project', sub_id=sp.id),
                    'export_excel_url': url_for('export_excel', sub_id=sp.id),
                    'export_zip_url': url_for('export_registration_forms', sub_id=sp.id)
                })
            results.append({
                'id': p.id,
                'code': p.code,
                'name': p.name,
                'sub_projects': sub_list
            })
        return jsonify({'projects': results})

    else:  # ✅ 客户视角
        now = datetime.now()

        if query:
            matched_projects = Project.query.filter(
                or_(
                    Project.code.contains(query),
                    Project.name.contains(query)
                )
            ).all()
        else:
            matched_projects = Project.query.all()

        results = []
        for project in matched_projects:
            # ✅ 改为判断项目时间
            if not (project.start_time and project.deadline):
                continue
            if not (project.start_time <= now <= project.deadline):
                continue

            subprojects = SubProject.query.filter_by(project_id=project.id).all()
            sub_list = []
            for sp in subprojects:
                sub_list.append({
                    "id": sp.id,
                    "segment_name": sp.segment_name or "默认标段"
                })

            results.append({
                "id": project.id,
                "code": project.code,
                "name": project.name,
                "has_segments": len(sub_list) > 0,
                "subprojects": sub_list
            })

        return jsonify({'projects': results})

# @app.route('/search_projects')
# def search_projects():
#     query = request.args.get("query", "").strip()
#
#     if session.get('admin'):  # ✅ 管理员视角
#         if query:
#             projects = Project.query.filter(
#                 or_(
#                     Project.code.contains(query),
#                     Project.name.contains(query)
#                 )
#             ).all()
#         else:
#             projects = Project.query.all()
#
#         results = []
#         for p in projects:
#             subprojects = SubProject.query.filter_by(project_id=p.id).all()
#             sub_list = []
#             for sp in subprojects:
#                 sub_list.append({
#                     'id': sp.id,
#                     'name': sp.segment_name or "默认标段",
#                     'start_time': sp.start_time.strftime('%Y年%m月%d日%H时%M分'),
#                     'deadline': sp.deadline.strftime('%Y年%m月%d日%H时%M分'),
#                     'deposit_amount': sp.deposit_amount,
#                     'file_path': sp.file_path,
#                     'view_bids_url': url_for('view_bids', project_id=p.id, sub_id=sp.id),
#                     'edit_project_url': url_for('edit_sub_project', sub_id=sp.id),
#                     'delete_project_url': url_for('delete_sub_project', sub_id=sp.id),
#                     'export_excel_url': url_for('export_excel', sub_id=sp.id),
#                     'export_zip_url': url_for('export_registration_forms', sub_id=sp.id)
#                 })
#             results.append({
#                 'id': p.id,
#                 'code': p.code,
#                 'name': p.name,
#                 'sub_projects': sub_list
#             })
#         return jsonify({'projects': results})
#
#     else:  # ✅ 客户视角（仅返回当前时间内进行中的标段）
#         if query:
#             matched_projects = Project.query.filter(
#                 or_(
#                     Project.code.contains(query),
#                     Project.name.contains(query)
#                 )
#             ).all()
#         else:
#             matched_projects = Project.query.all()
#
#         now = datetime.now()
#         results = []
#
#         for project in matched_projects:
#             ongoing_subprojects = SubProject.query.filter(
#                 SubProject.project_id == project.id,
#                 SubProject.start_time <= now,
#                 SubProject.deadline >= now
#             ).all()
#
#             if not ongoing_subprojects:
#                 continue
#
#             sub_list = []
#             for sp in ongoing_subprojects:
#                 sub_list.append({
#                     "id": sp.id,
#                     "name": sp.segment_name or "默认标段"
#                 })
#
#             results.append({
#                 "code": project.code,
#                 "name": project.name,
#                 "subprojects": sub_list
#             })
#
#         return jsonify({'projects': results})


# @app.route('/search_projects')
# def search_projects():
#     query = request.args.get("query")
#     if session.get('admin'):  # 管理员
#         if query:
#             # 管理员查询所有项目，不受截止时间限制
#             projects = Project.query.filter(
#                 or_(
#                     Project.code.contains(query),
#                     Project.name.contains(query)
#                 )
#             ).all()
#         else:
#             # 如果没有搜索条件，返回所有项目
#             projects = Project.query.all()
#
#         # 返回完整数据，包含操作按钮的链接
#         results = [{
#             'id': p.id,
#             'code': p.code,
#             'name': p.name,
#             'deadline': p.deadline.strftime('%Y年%m月%d日%H时%M分'),
#             'deposit_amount': p.deposit_amount,
#             'file_path': p.file_path,
#             'view_bids_url': url_for('view_bids', project_id=p.id),
#             'edit_project_url': url_for('edit_project', project_id=p.id),
#             'delete_project_url': url_for('delete_project', project_id=p.id)
#         } for p in projects]
#         return jsonify({'projects': results})
#     else:  # 客户视角：返回包含“进行中标段”的项目
#         if query:
#             matched_projects = Project.query.filter(
#                 or_(
#                     Project.code.contains(query),
#                     Project.name.contains(query)
#                 )
#             ).all()
#
#             now = datetime.now()
#             results = []
#
#             for project in matched_projects:
#                 ongoing_subprojects = SubProject.query.filter(
#                     SubProject.project_id == project.id,
#                     SubProject.start_time <= now,
#                     SubProject.deadline >= now
#                 ).all()
#
#                 if not ongoing_subprojects:
#                     continue  # 没有在进行中的标段就跳过这个项目
#
#                 subproject_list = []
#                 for sp in ongoing_subprojects:
#                     subproject_list.append({
#                         "id": sp.id,
#                         "name": sp.segment_name or "默认标段"
#                     })
#
#                 results.append({
#                     "code": project.code,
#                     "name": project.name,
#                     "subprojects": subproject_list
#                 })
#
#             return jsonify({'projects': results})

# ✅ 提供某个项目下的标段列表（客户端异步调用）
@app.route('/get_segments/<int:project_id>', methods=['GET'])
def get_segments(project_id):
    now = datetime.now()
    project = Project.query.get_or_404(project_id)
    if project.is_segmented:
        subprojects = SubProject.query.filter_by(project_id=project_id).all()
        result = []
        for sp in subprojects:
            result.append({
                'id': sp.id,
                'segment_name': sp.segment_name or '默认标段'
            })
        return jsonify({'subprojects': result})
    else:
        return jsonify({'subprojects': []})


# todo: 支持多标段
@app.route('/filter_projects')
def filter_projects():
    if not session.get('admin'):
        return jsonify({'projects': []})

    status_filter = request.args.get("deadline")
    leader_filter = request.args.get("leader", "").strip()
    now = datetime.now()

    query = Project.query

    if status_filter == "not_started":
        query = query.filter(Project.start_time > now)
    elif status_filter == "in_progress":
        query = query.filter(and_(Project.start_time <= now, Project.deadline >= now))
    elif status_filter == "ended":
        query = query.filter(Project.deadline < now)

    if leader_filter:
        matching_leaders = Leader.query.filter(Leader.name.contains(leader_filter)).all()
        emails = [l.email for l in matching_leaders]
        query = query.filter(Project.leader_email.in_(emails))

    projects = query.order_by(Project.deadline.desc()).all()

    results = []
    for p in projects:
        subprojects = SubProject.query.filter_by(project_id=p.id).all()
        sub_list = []
        for sp in subprojects:
            sub_list.append({
                'id': sp.id,
                'name': sp.segment_name or "默认标段",
                'start_time': sp.start_time.strftime('%Y年%m月%d日%H时%M分'),
                'deadline': sp.deadline.strftime('%Y年%m月%d日%H时%M分'),
                'deposit_amount': sp.deposit_amount,
                'file_path': sp.file_path,
                'view_bids_url': url_for('view_bids', project_id=p.id, sub_id=sp.id),
                'edit_project_url': url_for('edit_sub_project', sub_id=sp.id),
                'delete_project_url': url_for('delete_sub_project', sub_id=sp.id),
                'export_excel_url': url_for('export_excel', sub_id=sp.id),
                'export_zip_url': url_for('export_registration_forms', sub_id=sp.id)
            })
        results.append({
            'id': p.id,
            'code': p.code,
            'name': p.name,
            'sub_projects': sub_list
        })

    return jsonify({'projects': results})


# @app.route('/filter_projects')
# def filter_projects():
#     if not session.get('admin'):
#         return jsonify({'projects': []})
#
#     status_filter = request.args.get("deadline")  # 改名更准确
#     now = datetime.now()
#     leader_filter = request.args.get("leader")
#
#     query = Project.query
#
#     if status_filter == "not_started":
#         query = query.filter(Project.start_time > now)
#     elif status_filter == "in_progress":
#         query = query.filter(and_(Project.start_time <= now, Project.deadline >= now))
#     elif status_filter == "ended":
#         query = query.filter(Project.deadline < now)
#
#     if leader_filter:
#         matching_leaders = Leader.query.filter(Leader.name.contains(leader_filter)).all()
#         emails = [l.email for l in matching_leaders]
#         query = query.filter(Project.leader_email.in_(emails))
#
#     projects = query.order_by(Project.deadline.desc()).all()
#
#     results = [{
#         'id': p.id,
#         'code': p.code,
#         'name': p.name,
#         'deadline': p.deadline.strftime('%Y年%m月%d日%H时%M分'),
#         'deposit_amount': p.deposit_amount,
#         'file_path': p.file_path,
#         'view_bids_url': url_for('view_bids', project_id=p.id),
#         'edit_project_url': url_for('edit_project', project_id=p.id),
#         'delete_project_url': url_for('delete_project', project_id=p.id)
#     } for p in projects]
#
#     return jsonify({'projects': results})


# view_bids页面搜索公司名称模糊搜索
@app.route('/admin/search_supplier', methods=['GET'])
def search_supplier():
    if not session.get('admin'):
        return jsonify({'error': '未授权'}), 401

    keyword = request.args.get('q', '').strip()
    project_id = request.args.get('project_id', type=int)

    if not keyword or not project_id:
        return jsonify({'error': '参数缺失'}), 400

    results = Bid.query.filter(
        Bid.project_id == project_id,
        Bid.supplier_name.ilike(f"%{keyword}%")
    ).all()

    data = []
    for bid in results:
        data.append({
            'id': bid.id,
            'supplier_name': bid.supplier_name,
            'credit_code': bid.credit_code,
            'is_paid': bid.is_paid
        })

    return jsonify({'success': True, 'bids': data})


# 提交后不再回显
# ================== 返回修改路由 ==================
@app.route('/edit_last_bid', methods=['GET'])
def edit_last_bid():
    prefill_data = session.get('bid_form_data')
    if not prefill_data:
        flash("没有找到需要修改的数据", "danger")
        return redirect(url_for('index'))  # 如果没有数据，跳转回首页

    available_projects = Project.query.filter(Project.deadline > datetime.now()).all()
    return render_template('index.html', projects=available_projects, business_open=is_within_business_hours(),
                           prefill=prefill_data)


# ================== 成功页面：确认无误后清空表单 ==================
@app.route('/confirm_submission', methods=['POST'])
def confirm_submission():
    # 从 session 获取之前提交的表单数据
    prefill_data = session.get('bid_form_data')
    if not prefill_data:
        flash("❌ 无法确认投标信息，请先填写完整", "danger")
        return redirect(url_for('index'))

    # 获取当前投标记录并更新为确认状态
    project = Project.query.filter_by(code=prefill_data['project_code']).first()
    bid = Bid.query.filter_by(project_id=project.id, supplier_name=prefill_data['supplier_name']).first()

    if bid:
        bid.status = 'confirmed'
        db.session.commit()

    # 清空 session 中的表单数据
    session.pop('bid_form_data', None)
    # 跳转到首页
    return redirect(url_for('index'))


# 生成 Excel 统计表并下载。
# todo:支持标段
@app.route('/admin/project/<int:project_id>/export_excel', methods=['GET'])
def export_excel_project(project_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    project = Project.query.get_or_404(project_id)
    sub_projects = SubProject.query.filter_by(project_id=project.id).all()

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, Border, Side
    import os

    file_name = f"{project.name} 文件获取统计表.xlsx"
    file_path = os.path.join(Config.BASE_DIR, "static", "downloads", "Statistics", file_name)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    wb = Workbook()
    wb.remove(wb.active)  # 移除默认Sheet

    if sub_projects:
        # ✅ 多个标段，每个一个 sheet
        for sub in sub_projects:
            ws = wb.create_sheet(title=sub.segment_name)
            bids = Bid.query.filter(Bid.sub_project_id == sub.id, Bid.is_paid.in_([True, 1])).all()
            write_sheet(ws, project.code, sub.segment_name, bids, sub.deadline, sub.deposit_amount)
    else:
        # ✅ 无标段，导出项目本身
        ws = wb.create_sheet(title="项目登记表")
        bids = Bid.query.filter(Bid.project_id == project.id, Bid.is_paid.in_([True, 1])).all()
        write_sheet(ws, project.code, project.name, bids, project.deadline, project.deposit_amount)

    wb.save(file_path)
    upload_file_to_oss(file_path, f"统计表/{file_name}")

    return sf(
        file_path,
        as_attachment=True,
        download_name=file_name,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


def write_sheet(ws, project_code, title, bids, deadline, deposit_amount):
    ws.merge_cells('A1:F1')
    ws['A1'] = f"{project_code} - {title} 文件获取统计表"
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws['A1'].font = Font(size=14, bold=True)

    titles = ["序号", "单位名称", "委托代理人", "联系电话", "法人", "统一社会信用代码"]
    for col, text in enumerate(titles, 1):
        cell = ws.cell(row=2, column=col, value=text)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.font = Font(size=12)

    for idx, bid in enumerate(bids, start=1):
        ws.cell(row=2 + idx, column=1, value=idx)
        ws.cell(row=2 + idx, column=2, value=bid.supplier_name)
        ws.cell(row=2 + idx, column=3, value=bid.agent)
        ws.cell(row=2 + idx, column=4, value=bid.phone)
        ws.cell(row=2 + idx, column=5, value=bid.legal_person)
        ws.cell(row=2 + idx, column=6, value=bid.credit_code)
        for col in range(1, 7):
            cell = ws.cell(row=2 + idx, column=col)
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.font = Font(size=12)
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

    end_row = 2 + len(bids) + 1
    ws.merge_cells(f'A{end_row}:F{end_row}')
    ws[f'A{end_row}'] = (
        f"文件获取截止时间：{deadline.strftime('%Y年%m月%d日%H时%M分')}  "
        f"保证金金额：{round(deposit_amount, 2):.2f}元"
    )
    ws[f'A{end_row}'].alignment = Alignment(horizontal='center', vertical='center')
    ws[f'A{end_row}'].font = Font(size=12, bold=True)

    # 列宽设置
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 35
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 18
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 25

# @app.route('/admin/subproject/<int:sub_project_id>/export_excel', methods=['GET'])
# def export_excel(sub_project_id):
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#     subproject = SubProject.query.get_or_404(sub_project_id)
#     project = subproject.project  # 父项目
#     bids = Bid.query.filter(
#         Bid.sub_project_id == sub_project_id,
#         Bid.is_paid.in_([True, 1])
#     ).all()
#
#     file_name = f"{project.name} - {subproject.segment_name} 文件获取统计表.xlsx"
#     file_path = os.path.join(Config.BASE_DIR, "static", "downloads", "Statistics", file_name)
#     os.makedirs(os.path.dirname(file_path), exist_ok=True)
#
#     wb = Workbook()
#     ws = wb.active
#     ws.title = '投标信息'
#
#     # 表头
#     ws.merge_cells('A1:F1')
#     ws['A1'] = f"{project.code} - {subproject.segment_name} 文件获取统计表"
#     ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
#     ws['A1'].font = Font(size=14, bold=True)
#
#     titles = ["序号", "单位名称", "委托代理人", "联系电话", "法人", "统一社会信用代码"]
#     for col, title in enumerate(titles, start=1):
#         cell = ws.cell(row=2, column=col, value=title)
#         cell.alignment = Alignment(horizontal='center', vertical='center')
#         cell.font = Font(size=12)
#
#     # 内容
#     max_supplier_name_length = 0
#     for idx, bid in enumerate(bids, start=1):
#         ws.cell(row=2 + idx, column=1, value=idx)
#         ws.cell(row=2 + idx, column=2, value=bid.supplier_name)
#         ws.cell(row=2 + idx, column=3, value=bid.agent)
#         ws.cell(row=2 + idx, column=4, value=bid.phone)
#         ws.cell(row=2 + idx, column=5, value=bid.legal_person)
#         ws.cell(row=2 + idx, column=6, value=bid.credit_code)
#
#         if bid.supplier_name:
#             max_supplier_name_length = max(max_supplier_name_length, len(bid.supplier_name))
#
#         for col in range(1, 7):
#             cell = ws.cell(row=2 + idx, column=col)
#             cell.alignment = Alignment(horizontal='center', vertical='center')
#             cell.font = Font(size=12)
#             cell.border = Border(
#                 left=Side(style='thin'),
#                 right=Side(style='thin'),
#                 top=Side(style='thin'),
#                 bottom=Side(style='thin')
#             )
#
#     # 尾行：截止时间与保证金
#     end_row = 2 + len(bids) + 1
#     ws.merge_cells(f'A{end_row}:F{end_row}')
#     ws[f'A{end_row}'] = (
#         f"文件获取截止时间：{subproject.deadline.strftime('%Y年%m月%d日%H时%M分')}  "
#         f"保证金金额：{round(subproject.deposit_amount, 2):.2f}元"
#     )
#     ws[f'A{end_row}'].alignment = Alignment(horizontal='center', vertical='center')
#     ws[f'A{end_row}'].font = Font(size=12, bold=True)
#
#     # 列宽设置
#     ws.column_dimensions['A'].width = 8
#     base_width = 35
#     extra_width_per_char = 1.2
#     computed_width = max(base_width, min(base_width + (max_supplier_name_length - 10) * extra_width_per_char, 50))
#     ws.column_dimensions['B'].width = computed_width
#     ws.column_dimensions['C'].width = 20
#     ws.column_dimensions['D'].width = 18
#     ws.column_dimensions['E'].width = 15
#     ws.column_dimensions['F'].width = 25
#
#     wb.save(file_path)
#     upload_file_to_oss(file_path, f"统计表/{file_name}")
#
#     return sf(
#         file_path,
#         as_attachment=True,
#         download_name=file_name,
#         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
#     )


# @app.route('/admin/project/<int:project_id>/export_excel', methods=['GET'])
# def export_excel(project_id):
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#     project = Project.query.get_or_404(project_id)
#     # 交费了才有记录
#     bids = Bid.query.filter(Bid.project_id == project_id, Bid.is_paid.in_([True, 1])).all()
#
#     # if not bids:
#     #     flash("❌ 当前项目下没有已缴费客户，无法生成文件获取统计表", "danger")
#     #     return redirect(url_for('admin_panel'))
#
#     file_name = f"{project.name} {project.code} 文件获取统计表.xlsx"
#     file_path = os.path.join(Config.BASE_DIR, "static", "downloads", "Statistics", file_name)
#     # ✅ 确保目录存在
#     os.makedirs(os.path.dirname(file_path), exist_ok=True)
#     wb = Workbook()
#     ws = wb.active
#     ws.title = '投标信息'
#
#     # 合并第一行
#     ws.merge_cells('A1:F1')
#     ws['A1'] = f"{project.code} {project.name} 文件获取统计表"
#     ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
#     ws['A1'].font = Font(size=14, bold=True)
#
#     # 第二行列标题
#     titles = ["序号", "单位名称", "委托代理人", "联系电话", "法人", "统一社会信用代码"]
#     for col, title in enumerate(titles, start=1):
#         cell = ws.cell(row=2, column=col, value=title)
#         cell.alignment = Alignment(horizontal='center', vertical='center')
#         cell.font = Font(size=12)
#
#     # 填写客户数据
#     for idx, bid in enumerate(bids, start=1):
#         ws.cell(row=2 + idx, column=1, value=idx)  # 序号
#         ws.cell(row=2 + idx, column=2, value=bid.supplier_name)
#         ws.cell(row=2 + idx, column=3, value=bid.agent)
#         ws.cell(row=2 + idx, column=4, value=bid.phone)
#         ws.cell(row=2 + idx, column=5, value=bid.legal_person)
#         ws.cell(row=2 + idx, column=6, value=bid.credit_code)
#
#         # 居中、字体、边框
#         for col in range(1, 7):
#             cell = ws.cell(row=2 + idx, column=col)
#             cell.alignment = Alignment(horizontal='center', vertical='center')
#             cell.font = Font(size=12)
#             cell.border = Border(
#                 left=Side(style='thin'),
#                 right=Side(style='thin'),
#                 top=Side(style='thin'),
#                 bottom=Side(style='thin')
#             )
#
#     # 合并最后一行截止时间
#     end_row = 2 + len(bids) + 1
#     ws.merge_cells(f'A{end_row}:F{end_row}')
#     ws[
#         f'A{end_row}'] = f"文件获取截止时间：{project.deadline.strftime('%Y年%m月%d日%H时%M分')}  保证金金额：{round(project.deposit_amount, 2):.2f}元"
#     ws[f'A{end_row}'].alignment = Alignment(horizontal='center', vertical='center')
#     ws[f'A{end_row}'].font = Font(size=12, bold=True)
#
#     '''
#     要稍微在写入Excel之前，自己先扫描一下所有供应商名称的最大长度，
#     然后根据这个最大长度，动态调整列宽！
#     '''
#     # 动态计算单位名称列宽
#     max_supplier_name_length = 0
#     for bid in bids:
#         if bid.supplier_name:
#             max_supplier_name_length = max(max_supplier_name_length, len(bid.supplier_name))
#
#     # 设置列宽
#     ws.column_dimensions['A'].width = 8
#     ws.column_dimensions['B'].width = 35
#     ws.column_dimensions['C'].width = 20
#     ws.column_dimensions['D'].width = 18
#     ws.column_dimensions['E'].width = 15
#     ws.column_dimensions['F'].width = 25
#
#     # 单位名称列(B列)动态调整
#     base_width = 35
#     extra_width_per_char = 1.2
#     computed_width = max(base_width, min(base_width + (max_supplier_name_length - 10) * extra_width_per_char, 50))
#     ws.column_dimensions['B'].width = computed_width
#
#     wb.save(file_path)
#     upload_file_to_oss(file_path, f"统计表/{file_name}")
#     return sf(file_path, as_attachment=True, download_name=file_name,
#               mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# @app.route('/admin/project/<int:project_id>/export_excel', methods=['GET'])
#
# def export_excel(project_id):
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#     # project = Project.query.get_or_404(project_id)
#     # all_bids = Bid.query.filter_by(project_id=project_id).all()
#     #
#     # print(f"项目【{project.name}】下共找到 {len(all_bids)} 条Bid记录：")
#     # for bid in all_bids:
#     #     print(f"Bid ID: {bid.id}, 供应商名称: {bid.supplier_name}, 是否已缴费: {bid.is_paid}")
#     #     print({
#     #         "id": bid.id,
#     #         "supplier_name": bid.supplier_name,
#     #         "agent": bid.agent,
#     #         "phone": bid.phone,
#     #         "legal_person": bid.legal_person,
#     #         "credit_code": bid.credit_code,
#     #         "is_paid": bid.is_paid,
#     #         "status": bid.status
#     #     })
#
#     # 获取项目和投标信息
#     project = Project.query.get_or_404(project_id)
#     bids = Bid.query.filter(
#         Bid.project_id == project_id,
#         Bid.is_paid.in_([True, 1])  # ✅ 同时兼容True和1
#     ).all()
#
#     # 构建表格数据
#     data = []
#     for idx, bid in enumerate(bids, start=0):  # 序号从1开始
#         data.append([
#             idx,  # 序号
#             bid.supplier_name,  # 单位名称
#             bid.agent,  # 委托代理人
#             bid.phone,  # 联系电话
#             bid.legal_person,  # 法人
#             bid.credit_code,  # 统一社会信用代码
#         ])
#
#     # 创建一个 DataFrame
#     df = pd.DataFrame(data, columns=["序号", "单位名称", "委托代理人", "联系电话", "法人", "统一社会信用代码"])
#     # print(df)  # ✅ 打印检查
#     # 文件保存路径
#     file_name = f"{project.code} 文件获取统计表.xlsx"
#     file_path = os.path.join(Config.BASE_DIR, "static", "downloads", file_name)
#
#     # 创建 Excel 文件并保存
#     with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
#         df.to_excel(writer, index=False, sheet_name='投标信息')
#         workbook = writer.book
#         worksheet = workbook['投标信息']
#
#         # 合并第一行（项目编号 + 项目名称 + 文件获取统计表）
#         worksheet.merge_cells('A1:F1')
#         worksheet['A1'] = f"{project.code} {project.name} 文件获取统计表"
#         worksheet['A1'].alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')
#         # 设置第一行字体大小为 14
#         worksheet['A1'].font = Font(size=14, bold=True)
#
#         thin_border_first_row = Border(
#             left=Side(style=None),
#             right=Side(style=None),  # 不为右侧加边框
#             top=Side(style=None),
#             bottom=Side(style=None)
#         )
#
#         # 为第一行取消边框
#         for cell in worksheet["1"]:  # 第一行的所有单元格
#             cell.border = thin_border_first_row
#
#         # 第二行（属性名）
#         worksheet['A2'] = '序号'
#         worksheet['B2'] = '单位名称'
#         worksheet['C2'] = '委托代理人'
#         worksheet['D2'] = '联系电话'
#         worksheet['E2'] = '法人'
#         worksheet['F2'] = '统一社会信用代码'
#         worksheet['A2'].alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')
#         worksheet['B2'].alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')
#         worksheet['C2'].alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')
#         worksheet['D2'].alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')
#         worksheet['E2'].alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')
#         worksheet['F2'].alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')
#
#         # 第二行（属性名行）字体大小为 12
#         worksheet['A2'].font = Font(size=12)
#         worksheet['B2'].font = Font(size=12)
#         worksheet['C2'].font = Font(size=12)
#         worksheet['D2'].font = Font(size=12)
#         worksheet['E2'].font = Font(size=12)
#         worksheet['F2'].font = Font(size=12)
#
#         # 设置列宽
#         worksheet.column_dimensions['A'].width = 8  # 序号列
#         worksheet.column_dimensions['B'].width = 35  # 单位名称
#         worksheet.column_dimensions['C'].width = 20  # 委托代理人
#         worksheet.column_dimensions['D'].width = 18  # 联系电话
#         worksheet.column_dimensions['E'].width = 15  # 法人
#         worksheet.column_dimensions['F'].width = 25  # 统一社会信用代码
#
#         # 设置行高
#         worksheet.row_dimensions[1].height = 30  # 第一行
#         worksheet.row_dimensions[2].height = 25  # 第二行（属性名行）
#
#         # 为数据行设置边框
#         thin_border = Border(
#             left=Side(style='thin'),
#             right=Side(style='thin'),
#             top=Side(style='thin'),
#             bottom=Side(style='thin')
#         )
#
#         for cell in worksheet["2"]:  # 第二行的所有单元格
#             cell.border = thin_border
#
#             # 为数据行设置边框
#         for row in worksheet.iter_rows(min_row=3, max_row=2 + len(bids), min_col=1, max_col=6):
#             for cell in row:
#                 cell.border = thin_border
#                 cell.font = Font(size=12)
#                 cell.alignment = Alignment(horizontal='center', vertical='center')  # ✅ 居中对齐就靠这行
#
#             # 合并最后一行（文件获取截止时间 + 保证金金额），不加边框
#         end_row = str(len(bids) + 2)
#         worksheet.merge_cells(f'A{end_row}:F{end_row}')
#         worksheet[
#             f'A{end_row}'] = f"文件获取截止时间：{project.deadline.strftime('%Y年%m月%d日%H时%M分')}  保证金金额：{round(project.deposit_amount, 2):.2f}元"
#
#         # 设置居中对齐
#         worksheet[f'A{end_row}'].alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')
#
#         # ✅ 设置加粗字体
#         worksheet[f'A{end_row}'].font = Font(size=12, bold=True)
#         worksheet.row_dimensions[int(end_row)].height = 20  # 设置最后一行高度为20（数值）
#
#         # 不给最后一行添加边框
#         last_row_cells = worksheet["A" + str(len(bids) + 2) + ":F" + str(len(bids) + 2)][0]
#         for cell in last_row_cells:
#             cell.border = None
#
#     return sf(file_path, as_attachment=True, download_name=file_name,
#               mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
#

# todo：按子标段
# 导出客户登记表
def set_cell_width(cell, width_cm):
    cell_width_twips = int(width_cm * 567)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(cell_width_twips))
    tcPr.append(tcW)

def create_word_for_bid(project, sub, bid):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(14)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run('文件获取登记表')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(24)

    title_text = f"　项目名称：{project.name}"
    if sub:
        title_text += f" - {sub.segment_name}"
    else:
        title_text += "（无标段）"
    doc.add_paragraph(title_text)
    doc.add_paragraph(f"　项目编号：{project.code}")

    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.cell(0, 1).merge(table.cell(0, 3))
    table.cell(1, 1).merge(table.cell(1, 3))
    table.cell(3, 0).merge(table.cell(4, 0))
    table.cell(3, 1).merge(table.cell(4, 1))
    table.cell(5, 1).merge(table.cell(5, 3))
    table.cell(6, 1).merge(table.cell(6, 3))
    table.cell(7, 1).merge(table.cell(7, 3))

    table.cell(0, 0).text = "投标单位\n（供应商）名称"
    table.cell(0, 1).text = bid.supplier_name or ""

    table.cell(1, 0).text = "投标单位\n（供应商）地址"
    table.cell(1, 1).text = bid.supplier_address or ""

    table.cell(2, 0).text = "法定代表人"
    table.cell(2, 1).text = bid.legal_person or ""
    table.cell(2, 2).text = "统一社会信用代码"
    table.cell(2, 3).text = bid.credit_code or ""

    table.cell(3, 0).text = "委托代理人"
    table.cell(3, 1).text = bid.agent or ""
    table.cell(3, 2).text = "手机"
    table.cell(3, 3).text = bid.phone or ""

    table.cell(4, 2).text = "电子邮箱"
    table.cell(4, 3).text = bid.email or ""

    table.cell(5, 0).text = "招标（采购）\n文件"
    table.cell(5, 1).text = "☑已领取"

    table.cell(6, 0).text = "文件获取方式"
    method_text = "☑现场获取    ☐邮箱获取" if bid.file_method == "现场获取" else "☐现场获取    ☑邮箱获取"
    table.cell(6, 1).text = method_text

    table.cell(7, 0).text = "文件获取时间"
    file_time_display = bid.file_time or "     年       月       日       时       分"
    table.cell(7, 1).text = file_time_display

    widths_cm = [4.5, 4.5, 3.0, 5.0]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_cm[idx])
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    return doc

@app.route('/export_registration_forms_project/<int:project_id>')
def export_registration_forms_project(project_id):
    if not session.get('admin'):
        flash("❌ 无权限访问", "danger")
        return redirect(url_for('admin_panel'))

    project = Project.query.get_or_404(project_id)
    temp_dir = tempfile.mkdtemp()

    try:
        if not project.is_segmented:
            # ✅ 无标段：直接查 project_id 下的所有客户登记
            bids = Bid.query.filter_by(project_id=project_id).all()
            if not bids:
                flash("❌ 当前项目下没有客户登记信息", "danger")
                return redirect(url_for('admin_panel'))

            for bid in bids:
                doc = create_word_for_bid(project, None, bid)
                safe_name = ''.join(c for c in bid.supplier_name if c.isalnum() or c in "（）() ")
                doc.save(os.path.join(temp_dir, f"{safe_name} 文件获取登记表.docx"))

        else:
            # ✅ 有标段：每个标段一个文件夹，放入对应客户登记
            sub_projects = SubProject.query.filter_by(project_id=project_id).all()
            has_valid = False

            for sub in sub_projects:
                sub_bids = Bid.query.filter_by(sub_project_id=sub.id).all()
                if not sub_bids:
                    continue
                has_valid = True
                sub_folder = os.path.join(temp_dir, sub.segment_name)
                os.makedirs(sub_folder, exist_ok=True)
                for bid in sub_bids:
                    doc = create_word_for_bid(project, sub, bid)
                    safe_name = ''.join(c for c in bid.supplier_name if c.isalnum() or c in "（）() ")
                    doc.save(os.path.join(sub_folder, f"{safe_name} 文件获取登记表.docx"))

            if not has_valid:
                flash("❌ 所有标段下都没有客户登记信息", "danger")
                return redirect(url_for('admin_panel'))

        # ✅ 打包 zip
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for foldername, subfolders, filenames in os.walk(temp_dir):
                for filename in filenames:
                    abs_path = os.path.join(foldername, filename)
                    arcname = os.path.relpath(abs_path, temp_dir)
                    zipf.write(abs_path, arcname=arcname)
        zip_buffer.seek(0)

        # ✅ 文件保存 + 返回
        safe_project_name = ''.join(c for c in project.name if c.isalnum() or c in '（）() ')
        zip_name = f"{safe_project_name} 文件获取登记表.zip"

        save_dir = os.path.join(Config.BASE_DIR, "static", "downloads", "Registrations")
        os.makedirs(save_dir, exist_ok=True)
        save_path = os.path.join(save_dir, zip_name)
        with open(save_path, 'wb') as f:
            f.write(zip_buffer.getvalue())
            upload_file_to_oss(save_path, f"登记表/{zip_name}")

        encoded_filename = quote(zip_name)
        response = sf(zip_buffer,
                      mimetype='application/zip',
                      as_attachment=True,
                      download_name=zip_name)
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        return response

    finally:
        shutil.rmtree(temp_dir)



# @app.route('/export_registration_forms_project/<int:project_id>')
# def export_registration_forms_project(project_id):
#     if not session.get('admin'):
#         flash("❌ 无权限访问", "danger")
#         return redirect(url_for('admin_panel'))
#
#     project = Project.query.get_or_404(project_id)
#     bids = Bid.query.filter_by(project_id=project_id).all()
#
#     if not bids:
#         flash("❌ 当前项目下没有客户登记信息", "danger")
#         return redirect(url_for('admin_panel'))
#
#     temp_dir = tempfile.mkdtemp()
#
#     try:
#         for bid in bids:
#             doc = Document()
#
#             # 设置字体样式
#             style = doc.styles['Normal']
#             style.font.name = '宋体'
#             style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#             style.font.size = Pt(14)
#
#             # 添加标题
#             title = doc.add_paragraph()
#             title.alignment = 1
#             run = title.add_run('文件获取登记表')
#             run.font.name = '宋体'
#             run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#             run.font.size = Pt(24)
#
#             # 添加项目信息
#             doc.add_paragraph(f"　项目名称：{project.name}（无标段）")
#             doc.add_paragraph(f"　项目编号：{project.code}")
#
#             # 创建表格
#             table = doc.add_table(rows=8, cols=4)
#             table.style = 'Table Grid'
#             table.alignment = WD_TABLE_ALIGNMENT.CENTER
#
#             table.cell(0, 1).merge(table.cell(0, 3))
#             table.cell(1, 1).merge(table.cell(1, 3))
#             table.cell(3, 0).merge(table.cell(4, 0))
#             table.cell(3, 1).merge(table.cell(4, 1))
#             table.cell(5, 1).merge(table.cell(5, 3))
#             table.cell(6, 1).merge(table.cell(6, 3))
#             table.cell(7, 1).merge(table.cell(7, 3))
#
#             table.cell(0, 0).text = "投标单位\n（供应商）名称"
#             table.cell(0, 1).text = bid.supplier_name or ""
#
#             table.cell(1, 0).text = "投标单位\n（供应商）地址"
#             table.cell(1, 1).text = bid.supplier_address or ""
#
#             table.cell(2, 0).text = "法定代表人"
#             table.cell(2, 1).text = bid.legal_person or ""
#             table.cell(2, 2).text = "统一社会信用代码"
#             table.cell(2, 3).text = bid.credit_code or ""
#
#             table.cell(3, 0).text = "委托代理人"
#             table.cell(3, 1).text = bid.agent or ""
#             table.cell(3, 2).text = "手机"
#             table.cell(3, 3).text = bid.phone or ""
#
#             table.cell(4, 2).text = "电子邮箱"
#             table.cell(4, 3).text = bid.email or ""
#
#             table.cell(5, 0).text = "招标（采购）\n文件"
#             table.cell(5, 1).text = "☑已领取"
#
#             table.cell(6, 0).text = "文件获取方式"
#             method_text = "☑现场获取    ☐邮箱获取" if bid.file_method == "现场获取" else "☐现场获取    ☑邮箱获取"
#             table.cell(6, 1).text = method_text
#
#             table.cell(7, 0).text = "文件获取时间"
#             file_time_display = bid.file_time or "     年       月       日       时       分"
#             table.cell(7, 1).text = file_time_display
#
#             # 设置列宽
#             widths_cm = [4.5, 4.5, 3.0, 5.0]
#             for row in table.rows:
#                 for idx, cell in enumerate(row.cells):
#                     set_cell_width(cell, widths_cm[idx])
#
#             # 设置行高和居中
#             for idx, row in enumerate(table.rows):
#                 row.height_rule = 1  # EXACT
#                 if idx in [2, 3, 4]:
#                     row.height = Pt(50)
#                 elif idx in [5, 6, 7]:
#                     row.height = Pt(60)
#                 else:
#                     row.height = Pt(80)
#                 for cell in row.cells:
#                     cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
#                     for paragraph in cell.paragraphs:
#                         paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#             safe_supplier_name = ''.join(c for c in bid.supplier_name if c.isalnum() or c in "（）() ")
#             word_filename = f"{safe_supplier_name} 文件获取登记表.docx"
#             word_path = os.path.join(temp_dir, word_filename)
#             doc.save(word_path)
#             del doc
#
#         # 打包为 zip，文件名也不包含标段名
#         zip_buffer = BytesIO()
#         with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
#             for filename in os.listdir(temp_dir):
#                 file_path = os.path.join(temp_dir, filename)
#                 zipf.write(file_path, arcname=filename)
#         zip_buffer.seek(0)
#
#         safe_project_name = ''.join(c for c in project.name if c.isalnum() or c in '（）() ')
#         safe_filename = f"{safe_project_name} 文件获取登记表.zip"
#
#         save_dir = os.path.join(Config.BASE_DIR, "static", "downloads", "Registrations")
#         os.makedirs(save_dir, exist_ok=True)
#         save_path = os.path.join(save_dir, safe_filename)
#         with open(save_path, 'wb') as f:
#             f.write(zip_buffer.getvalue())
#             upload_file_to_oss(save_path, f"登记表/{safe_filename}")
#
#         encoded_filename = quote(safe_filename)
#         response = sf(zip_buffer,
#                       mimetype='application/zip',
#                       as_attachment=True,
#                       download_name=safe_filename)
#         response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
#         return response
#
#     finally:
#         shutil.rmtree(temp_dir)


# @app.route('/export_registration_forms/<int:project_id>')
# def export_registration_forms(project_id):
#     if not session.get('admin'):
#         flash("❌ 无权限访问", "danger")
#         return redirect(url_for('admin_panel'))
#
#     project = Project.query.get_or_404(project_id)
#     bids = Bid.query.filter_by(project_id=project_id).all()
#
#     if not bids:
#         flash("❌ 当前项目下没有客户登记信息", "danger")
#         return redirect(url_for('admin_panel'))
#
#     temp_dir = tempfile.mkdtemp()
#
#     try:
#         for bid in bids:
#             doc = Document()
#
#             # 字体
#             style = doc.styles['Normal']
#             style.font.name = '宋体'
#             style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#             style.font.size = Pt(14)  # 四号字体，大一点
#
#             # 标题
#             title = doc.add_paragraph()
#             title.alignment = 1
#             run = title.add_run('文件获取登记表')
#             run.font.name = '宋体'
#             run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
#             run.font.size = Pt(24)  # 小一号，标题
#             run.bold = False  # 加粗吗
#
#             doc.add_paragraph(f"　项目名称：{project.name}")  # 用全角空格（U+3000）
#             doc.add_paragraph(f"　项目编号：{project.code}")
#             # 表格
#             table = doc.add_table(rows=8, cols=4)
#             table.style = 'Table Grid'
#             table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中
#
#             # 合并单元格
#             table.cell(0, 1).merge(table.cell(0, 3))
#             table.cell(1, 1).merge(table.cell(1, 3))
#             table.cell(3, 0).merge(table.cell(4, 0))
#             table.cell(3, 1).merge(table.cell(4, 1))
#             table.cell(5, 1).merge(table.cell(5, 3))
#             table.cell(6, 1).merge(table.cell(6, 3))
#             table.cell(7, 1).merge(table.cell(7, 3))
#
#             # 填内容
#             table.cell(0, 0).text = "投标单位\n（供应商）名称"
#             table.cell(0, 1).text = bid.supplier_name or ""
#
#             table.cell(1, 0).text = "投标单位\n（供应商）地址"
#             table.cell(1, 1).text = bid.supplier_address or ""
#
#             table.cell(2, 0).text = "法定代表人"
#             table.cell(2, 1).text = bid.legal_person or ""
#             table.cell(2, 2).text = "统一社会信用代码"
#             table.cell(2, 3).text = bid.credit_code or ""
#
#             table.cell(3, 0).text = "委托代理人"
#             table.cell(3, 1).text = bid.agent or ""
#             table.cell(3, 2).text = "手机"
#             table.cell(3, 3).text = bid.phone or ""
#
#             table.cell(4, 2).text = "电子邮箱"
#             table.cell(4, 3).text = bid.email or ""
#
#             table.cell(5, 0).text = "招标（采购）\n文件"
#             table.cell(5, 1).text = "☑已领取"
#
#             table.cell(6, 0).text = "文件获取方式"
#             method_text = "☑现场获取    ☐邮箱获取" if bid.file_method == "现场获取" else "☐现场获取    ☑邮箱获取"
#             table.cell(6, 1).text = method_text
#
#             table.cell(7, 0).text = "文件获取时间"
#             file_time_display = bid.file_time if bid.file_time else "     年       月       日       时       分"
#             table.cell(7, 1).text = file_time_display
#
#             # ✅ 设置列宽
#             widths_cm = [4.5, 4.5, 3.0, 5.0]  # 调小右边
#             for row in table.rows:
#                 for idx, cell in enumerate(row.cells):
#                     set_cell_width(cell, widths_cm[idx])
#
#             # ✅ 设置居中+拉高行高
#             for idx, row in enumerate(table.rows):
#                 # if idx in [0, 1]:  # 第0行、第1行：允许自动高度
#                 #     row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
#                 #     # ✅ 重要：手动给每个cell里打一个空行，撑开起步高度
#                 #     for cell in row.cells:
#                 #         if not cell.text.strip():  # 如果本来是空的单元格
#                 #             para = cell.paragraphs[0]
#                 #             para.add_run("\n\n\n\n")  # 加个换行撑高度
#                 # else:  # 其他行保持固定高度
#                 row.height_rule = 1  # EXACT
#                 if idx == 2 or idx == 3 or idx == 4:
#                     row.height = Pt(50)  # ✅ 法定代表人和委托代理人这三行稍微矮一点（比如50pt）
#                 elif idx == 5 or idx == 6 or idx == 7:
#                     row.height = Pt(60)  # ✅ 最后三行，招标文件/获取方式/获取时间，也稍矮
#                 else:
#                     row.height = Pt(80)  # ✅ 其他行继续80pt高
#                 for cell in row.cells:
#                     cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
#                     for paragraph in cell.paragraphs:
#                         paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#             # 保存文件
#             safe_supplier_name = ''.join(c for c in bid.supplier_name if c.isalnum() or c in "（）() ")
#             word_filename = f"{safe_supplier_name} 文件获取登记表.docx"
#             word_path = os.path.join(temp_dir, word_filename)
#             doc.save(word_path)
#             del doc  # ✅ 强制销毁doc对象
#         # 打包所有 Word 文件为 zip
#         zip_buffer = BytesIO()
#         with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
#             for filename in os.listdir(temp_dir):
#                 file_path = os.path.join(temp_dir, filename)
#                 zipf.write(file_path, arcname=filename)
#
#         zip_buffer.seek(0)
#
#         # 安全处理文件名
#         safe_project_name = ''.join(c for c in project.name if c.isalnum() or c in '（）() ')
#         safe_filename = f"{safe_project_name} {project.code} 文件获取登记表.zip"
#
#         # 先准备保存目录
#         save_dir = os.path.join(Config.BASE_DIR, "static", "downloads", "Registrations")
#         os.makedirs(save_dir, exist_ok=True)  # ✅ 确保目录存在
#
#         # 保存一份zip到本地
#         save_path = os.path.join(save_dir, safe_filename)
#         with open(save_path, 'wb') as f:
#             f.write(zip_buffer.getvalue())
#             upload_file_to_oss(save_path, f"登记表/{safe_filename}")
#
#         # 注意：保存完不要seek(0)，因为zip_buffer已经是0了
#
#         # print(f"生成的safe_project_name是：{safe_project_name}")
#         # print(f"生成的safe_filename是：{safe_filename}")
#
#         encoded_filename = quote(safe_filename)
#         response = sf(
#             zip_buffer,
#             mimetype='application/zip',
#             as_attachment=True,
#             download_name=safe_filename  # ✅ 这里必须加
#         )
#         response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
#         return response
#
#     finally:
#         shutil.rmtree(temp_dir)


# DeepSeek API 信息
DEEPSEEK_API_KEY = "sk-ecba15c732384488836197dda0e81ac0"
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
USE_PROXY = False
# 👇 更明确、更贴合投标业务的 system prompt
SYSTEM_PROMPT = {
    "role": "system",
    "content": (
        "你是云南国合建设招标咨询公司的AI智能问答助手，负责为客户提供详尽、准确、实用的帮助,请用专业、友好的语气尽量详细回答客户提出的问题。"
        "请以亲切、专业、主动的语气尽可能详细地回答用户提出的问题，尤其是与报名流程、标书购买、"
        "缴费方式、保证金处理、开标时间、材料准备、发票问题、对公对私支付相关的问题。"
        "在无法确定答案时，也请你主动推测最合理的回复，而不是直接拒绝回答或说“我不知道”。"
        "如果用户的问题实在与你职责无关，请礼貌地说“该问题不在我的职责范围内”。"
        "你的目标是帮助客户解决他们的一切疑问，成为他们信赖的投标助手。"
        "你特别擅长解释报名、缴费、发票、保证金、投标材料、开标时间等事务，必要时可以主动推测合理回答。"
        "请避免说“我不知道”或“我无法回答”，除非问题明显不合理。"
        "如果你不知道，请合理推测!!!!"
    )
}


@app.route('/ask_deepseek', methods=['POST'])
def ask_deepseek():
    user_question = request.json.get("question", "").strip()
    if not user_question:
        return jsonify({"error": "问题不能为空"}), 400

    # ✅ 初始化聊天记录
    if "chat_history" not in session:
        session["chat_history"] = [SYSTEM_PROMPT]

    # ✅ 添加用户提问到历史中
    session["chat_history"].append({
        "role": "user",
        "content": f"请模拟一位标书服务人员回答以下问题：{user_question}"
    })
    # # ✅ 添加用户提问到历史中
    # session["chat_history"].append({"role": "user", "content": user_question})
    # ✅ 构造请求 payload
    payload = {
        "model": "deepseek-chat",
        "messages": session["chat_history"],
        "temperature": 0.3,  # 控制回答更丰富
        "top_p": 1.0  # 保证开放性
        # 有些平台可以添加 "max_tokens": 1024 以支持更长回答

    }

    # ✅ 设置代理配置（如需）
    proxies = {
        "http": "http://127.0.0.1:7890",
        "https": "http://127.0.0.1:7890"
    } if USE_PROXY else {}

    try:
        response = requests.post(
            DEEPSEEK_API_URL,
            headers={
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "Content-Type": "application/json"
            },
            json=payload,
            timeout=15,
            proxies=proxies
        )

        result = response.json()
        if "choices" in result and len(result["choices"]) > 0:
            reply = result["choices"][0]["message"]["content"]
            # ✅ 追加 AI 回复到上下文
            session["chat_history"].append({"role": "assistant", "content": reply})
            return jsonify({"answer": reply})
        else:
            return jsonify({"error": "DeepSeek 无回复"}), 500

    except Exception as e:
        return jsonify({"error": f"请求失败：{str(e)}"}), 500


# 清空机器人回答内容
@app.route('/reset_chat', methods=['POST'])
def reset_chat():
    session.pop("chat_history", None)
    return jsonify({"status": "reset"})


# 阿里云网盘

# @app.route('/admin/backup_to_oss', methods=['POST'])
# def backup_to_oss():
#     # 备份数据库到 OSS
#     db_path = os.path.join(Config.BASE_DIR, "data", "sqlite.db")
#     remote_db_path = f"数据库备份/sqlite_{datetime.now().strftime('%Y%m%d_%H%M')}.db"
#     # upload_file_to_oss(db_path, remote_db_path)
#
#     # 备份 static/uploads 下的所有文件到 OSS
#     for root, _, files in os.walk(Config.UPLOAD_FOLDER):
#         for file in files:
#             local_path = os.path.join(root, file)
#             relative_path = os.path.relpath(local_path, Config.UPLOAD_FOLDER)
#             oss_path = f"项目附件/{relative_path.replace(os.sep, '/')}"
#             upload_file_to_oss(local_path, oss_path)
#
#     return jsonify({"message": "✅ 已成功将数据库与文件备份至 OSS 云端！"})


# 手动触发 OSS 备份”
@app.route('/admin/manual_backup_to_oss', methods=['POST'])
def manual_backup_to_oss():
    from sync_oss import upload_db_to_oss, sync_static_to_oss

    upload_db_to_oss()
    sync_static_to_oss()

    response = jsonify({"message": "✅ 已成功将数据库与文件备份至 OSS 云端！"})
    response.headers['Content-Type'] = 'application/json'
    return response


@app.route('/admin/clear_local_data', methods=['POST'])
def clear_local_data():
    # 清空数据库表
    db.session.query(Bid).delete()
    db.session.query(Project).delete()
    db.session.query(MailLog).delete()
    db.session.commit()

    # 清空 uploads 和 temp_uploads 目录
    for folder in [Config.UPLOAD_FOLDER, Config.BDP_UPLOAD_TEMP_DIR]:
        if os.path.exists(folder):
            shutil.rmtree(folder)
            os.makedirs(folder)

    return jsonify({"message": "✅ 本地业务数据已清除！"})


# 自动清理超过 30天的临时 JSON 文件
@app.route('/admin/cleanup_temp_uploads', methods=['POST'])
def cleanup_temp_uploads():
    folder = Config.BDP_UPLOAD_TEMP_DIR  # temp_uploads
    now = time.time()
    cutoff = now - 30 * 86400  # 30天

    if not os.path.exists(folder):
        return jsonify({"message": "❌ 目录不存在，无需清理"})

    deleted_files = []
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        if os.path.isfile(file_path):
            file_mtime = os.path.getmtime(file_path)
            if file_mtime < cutoff:
                os.remove(file_path)
                deleted_files.append(filename)

    return jsonify({"message": f"✅ 清理完成，共删除 {len(deleted_files)} 个文件"})


# todo:支持按标段（SubProject）维度发送日报邮件（每个标段单独统计、单独发送）
def send_daily_project_summary():
    print("🕔 [定时任务] 开始发送项目日报邮件...")

    with app.app_context():
        now = datetime.now()
        today_str = now.strftime('%Y年%m月%d日%H时%M分')
        tolerate_seconds = 60  # 容忍1分钟小延迟

        sub_projects = SubProject.query.all()

        for sub in sub_projects:
            project = sub.project  # 父项目
            if not project or not project.leader_email:
                continue

            # ⏰ 判断当前时间是否仍在标段文件获取阶段
            if not (sub.start_time <= now <= sub.deadline + timedelta(seconds=tolerate_seconds)):
                print(f"⏳ 标段 [{sub.segment_name}] 不在发送时间范围内，跳过")
                continue

            # 查询该标段的已缴费投标单位
            bids = Bid.query.filter_by(sub_project_id=sub.id, is_paid=True).all()
            company_names = [bid.supplier_name for bid in bids]
            company_count = len(company_names)
            companies_text = "、".join(company_names)

            body = (
                f"尊敬的领导：{project.code}（{project.name}）下标段「{sub.segment_name}」截至{today_str}，"
                f"{'共' if company_count > 0 else ''}{companies_text if companies_text else ''}"
                f"{company_count}家单位获取文件；本标段文件获取截止时间为{sub.deadline.strftime('%Y年%m月%d日%H时%M分')}，特汇报！"
            )

            # ✅ 群发给项目组长 + 成员
            recipient_emails = set([project.leader_email] + [m.email for m in project.members])

            for email in recipient_emails:
                print(f"🚀 正在发送标段「{sub.segment_name}」日报给: {email}")
                result = send_email_with_attachment(
                    recipient_email=email,
                    subject=f"【{project.code}】【{sub.segment_name}】标段日报通知",
                    body=body,
                    base_filename=None
                )

                # ✅ 写入日志
                log_entry = MailLog(
                    project_id=project.id,
                    status="success" if result['status'] == 'success' else "failed",
                    message=f"发送{'成功' if result['status'] == 'success' else '失败'}: {email}（{sub.segment_name}）"
                )
                db.session.add(log_entry)
                db.session.commit()

                # ⏳ 失败时补发一次
                if result['status'] != 'success':
                    print(f"⚠️ 第一次发送失败，准备补发给 {email} ...")
                    retry_result = send_email_with_attachment(
                        recipient_email=email,
                        subject=f"【{project.code}】【{sub.segment_name}】标段日报通知（补发）",
                        body=body,
                        base_filename=None
                    )

                    retry_log = MailLog(
                        project_id=project.id,
                        status="success" if retry_result['status'] == 'success' else "failed",
                        message=f"补发{'成功' if retry_result['status'] == 'success' else '失败'}: {email}（{sub.segment_name}）"
                    )
                    db.session.add(retry_log)
                    db.session.commit()


# def send_daily_project_summary():
#     print("🕔 [定时任务] 开始发送项目日报邮件...")
#
#     with app.app_context():
#         now = datetime.now()
#         today_str = now.strftime('%Y年%m月%d日%H时%M分')
#
#         tolerate_seconds = 60  # 容忍1分钟的小延迟
#         projects = Project.query.all()
#
#         for project in projects:
#             if not project.leader_email:
#                 continue  # 没有组长邮箱就跳过
#
#             # 🛑 加上起始截止时间判断
#             project_start = project.start_time
#             project_deadline = project.deadline
#
#             if not (project_start <= now <= project_deadline + timedelta(seconds=tolerate_seconds)):
#                 print(f"⏳ 项目 {project.code} {project.name} 不在发送时间范围内，不再发送。")
#                 continue  # 超过截止+1分钟，就不再发
#
#             bids = Bid.query.filter_by(project_id=project.id, is_paid=True).all()
#             company_names = [bid.supplier_name for bid in bids]
#             company_count = len(company_names)
#             companies_text = "、".join(company_names)
#
#             body = (
#                 f"尊敬的领导：{project.code} {project.name}至{today_str}，"
#                 f"共{companies_text if companies_text else ''}"
#                 f"{company_count}家单位获取文件；本项目文件获取截止时间为{project.deadline.strftime('%Y年%m月%d日%H时%M分')}，特汇报！"
#             )
#
#             # ✅ 群发给组长和所有成员
#             recipient_emails = set([project.leader_email] + [m.email for m in project.members])
#
#             for email in recipient_emails:
#                 print(f"🚀 正在发送日报邮件给: {email}")
#                 result = send_email_with_attachment(
#                     recipient_email=email,
#                     subject=f"【{project.code}】项目日报通知",
#                     body=body,
#                     base_filename=None
#                 )
#
#                 if result['status'] == 'success':
#                     db.session.add(MailLog(
#                         project_id=project.id,
#                         status="success",
#                         message=f"发送成功: {email}"
#                     ))
#                     db.session.commit()
#                 else:
#                     db.session.add(MailLog(
#                         project_id=project.id,
#                         status="failed",
#                         message=f"第一次发送失败: {email} 错误信息: {result['error']}"
#                     ))
#                     db.session.commit()
#
#                     print(f"⚠️ 第一次发送失败，准备补发给 {email} ...")
#                     retry_result = send_email_with_attachment(
#                         recipient_email=email,
#                         subject=f"【{project.code}】项目日报通知（补发）",
#                         body=body,
#                         base_filename=None
#                     )
#
#                     if retry_result['status'] == 'success':
#                         db.session.add(MailLog(
#                             project_id=project.id,
#                             status="success",
#                             message=f"补发成功: {email}"
#                         ))
#                     else:
#                         db.session.add(MailLog(
#                             project_id=project.id,
#                             status="failed",
#                             message=f"补发失败: {email} 错误信息: {retry_result['error']}"
#                         ))
#                     db.session.commit()
#
#

# 项目邮件发送记录页面
# todo:支持标发送
@app.route('/admin/mail_logs', methods=['GET', 'POST'])
def view_mail_logs():
    if not session.get('admin'):
        return redirect(url_for('login'))

    start_time = request.form.get('start_time')
    end_time = request.form.get('end_time')

    query = MailLog.query

    if start_time and end_time:
        try:
            start_dt = datetime.strptime(start_time, '%Y-%m-%dT%H:%M')
            end_dt = datetime.strptime(end_time, '%Y-%m-%dT%H:%M')
            query = query.filter(MailLog.sent_at >= start_dt, MailLog.sent_at <= end_dt)
        except:
            query = query.filter(False)  # 空结果

    logs = query.order_by(MailLog.sent_at.desc()).limit(200).all()

    # 🧠 提前加载关联的子标段和项目，避免前端取不到 project.name 等
    for log in logs:
        if log.sub_project:
            _ = log.sub_project.project  # 提前访问，防止懒加载失败

    return render_template('view_mail_logs.html', logs=logs)


# @app.route('/admin/mail_logs', methods=['GET', 'POST'])
# def view_mail_logs():
#     if not session.get('admin'):
#         return redirect(url_for('login'))
#
#     start_time = request.form.get('start_time')
#     end_time = request.form.get('end_time')
#
#     if start_time and end_time:
#         # 用户提交了筛选时间
#         try:
#             start_dt = datetime.strptime(start_time, '%Y-%m-%dT%H:%M')
#             end_dt = datetime.strptime(end_time, '%Y-%m-%dT%H:%M')
#             logs = MailLog.query.filter(
#                 MailLog.sent_at >= start_dt,
#                 MailLog.sent_at <= end_dt
#             ).order_by(MailLog.sent_at.desc()).limit(200).all()
#         except:
#             logs = []
#     else:
#         # 默认加载最新40条
#         logs = MailLog.query.order_by(MailLog.sent_at.desc()).limit(40).all()
#
#     return render_template('view_mail_logs.html', logs=logs)


@app.route('/admin/retry_mail/<int:log_id>', methods=['POST'])
def retry_mail(log_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    log = MailLog.query.get_or_404(log_id)

    # 找回项目
    project = Project.query.get_or_404(log.project_id)
    if not project.leader_email:
        flash("❌ 找不到项目组长邮箱，无法重发！", "danger")
        return redirect(url_for('view_mail_logs'))

    # 重新组装日报内容
    now = datetime.now()
    today_str = now.strftime('%Y年%m月%d日%H时%M分')

    bids = Bid.query.filter_by(project_id=project.id, is_paid=True).all()
    company_names = [bid.supplier_name for bid in bids]
    company_count = len(company_names)
    companies_text = "、".join(company_names)

    body = (
        f"尊敬的领导：{project.code} {project.name}至{today_str}，"
        f"共{companies_text if companies_text else ''}"
        f"{company_count}家单位获取文件；本项目文件获取截止时间为{project.deadline.strftime('%Y年%m月%d日%H时%M分')}，特汇报！"
    )

    # 发送
    result = send_email_with_attachment(
        recipient_email=project.leader_email,
        subject=f"【{project.code}】项目日报通知（补发）",
        body=body
    )

    if result['status'] == 'success':
        flash("✅ 补发邮件成功！", "success")
        # 可以记录新的MailLog成功记录
    else:
        flash("❌ 补发邮件失败：" + result['error'], "danger")

    return redirect(url_for('view_mail_logs'))


import pandas as pd


@app.route('/admin/export_mail_logs')
def export_mail_logs():
    if not session.get('admin'):
        return redirect(url_for('login'))

    logs = MailLog.query.order_by(MailLog.sent_at.desc()).all()

    data = []
    for idx, log in enumerate(logs, 1):
        data.append({
            '序号': idx,
            '发送时间': log.sent_at.strftime('%Y-%m-%d %H:%M:%S') if log.sent_at else '',
            '项目名称': log.project.name if log.project else '无关联项目',
            '项目编号': log.project.code if log.project else '无编号',
            '状态': '成功' if log.status == 'success' else '失败',
            '信息': log.message or '无',
        })

    df = pd.DataFrame(data)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'邮件发送记录_{timestamp}.xlsx'

    # 保存到服务器
    save_dir = os.path.join('static', 'downloads', 'Mail_logs')
    os.makedirs(save_dir, exist_ok=True)
    save_path = os.path.join(save_dir, filename)
    df.to_excel(save_path, index=False, sheet_name='MailLogs', engine='openpyxl')

    # 返回给浏览器
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='MailLogs')
    output.seek(0)

    return sf(output, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.route('/api/leaders')
def get_leaders():
    keyword = request.args.get("q", "")
    if keyword:
        leaders = Leader.query.filter(Leader.name.contains(keyword)).all()
    else:
        leaders = Leader.query.all()

    return jsonify([
        {"id": leader.id, "name": leader.name, "email": leader.email}  # ✅ 加上 id
        for leader in leaders
    ])


# 进入组长管理页面
@app.route('/admin/leaders')
def manage_leaders():
    if not session.get('admin'):
        return redirect(url_for('login'))
    leaders = Leader.query.order_by(Leader.id).all()
    return render_template('manage_leaders.html', leaders=leaders)


# 添加新组长
@app.route('/admin/leaders/add', methods=['POST'])
def add_leader():
    if not session.get('admin'):
        return redirect(url_for('login'))
    name = request.form['name'].strip()
    email = request.form['email'].strip()

    # ✅ 邮箱格式验证
    if not re.match(r'^[^@]+@[^@]+\.(com|cn|net)$', email):
        flash("❌ 邮箱格式不正确，请重新填写", "danger")
        return redirect(url_for('manage_leaders'))

    if Leader.query.filter_by(name=name).first():
        flash("❌ 已存在同名项目组长", "danger")
    else:
        db.session.add(Leader(name=name, email=email))
        db.session.commit()
        flash("✅ 添加成功", "success")
    return redirect(url_for('manage_leaders'))


# 删除组长 只可删除没项目或项目已经结束的
@app.route('/admin/leaders/delete/<int:leader_id>', methods=['POST'])
def delete_leader(leader_id):
    if not session.get('admin'):
        return redirect(url_for('login'))
    leader = Leader.query.get_or_404(leader_id)

    # ✅ 查询该组长负责的项目是否存在未结束的
    active_projects = Project.query.filter(
        Project.leader_email == leader.email,
        Project.deadline > datetime.now()
    ).all()

    if active_projects:
        flash("❌ 该组长仍有未结束项目，无法删除", "danger")
        return redirect(url_for('manage_leaders'))

    db.session.delete(leader)
    db.session.commit()
    flash("🗑 项目组长已删除", "success")
    return redirect(url_for('manage_leaders'))


# ✅ 进入编辑页面
@app.route('/admin/leaders/edit/<int:leader_id>', methods=['GET', 'POST'])
def edit_leader(leader_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    leader = Leader.query.get_or_404(leader_id)

    if request.method == 'POST':
        new_name = request.form['name'].strip()
        new_email = request.form['email'].strip()

        # ✅ 邮箱格式验证
        if not re.match(r'^[^@]+@[^@]+\.(com|cn|net)$', new_email):
            flash("❌ 邮箱格式不正确", "danger")
            return redirect(url_for('edit_leader', leader_id=leader_id))

        # ✅ 同名校验（除自己外）
        conflict = Leader.query.filter(Leader.name == new_name, Leader.id != leader.id).first()
        if conflict:
            flash("❌ 已存在同名组长", "danger")
            return redirect(url_for('edit_leader', leader_id=leader_id))

        # ✅ 更新
        leader.name = new_name
        leader.email = new_email
        db.session.commit()
        flash("✅ 修改成功", "success")
        return redirect(url_for('manage_leaders'))

    return render_template("edit_leader.html", leader=leader)


@app.route('/api/leaders_by_ids')
def get_leaders_by_ids():
    ids = request.args.get('ids', '')
    try:
        id_list = [int(i) for i in ids.split(',') if i.strip().isdigit()]
    except ValueError:
        return jsonify([])

    leaders = Leader.query.filter(Leader.id.in_(id_list)).all()
    return jsonify([
        {'id': l.id, 'name': l.name, 'email': l.email}
        for l in leaders
    ])


@app.route('/admin/cleanup_uploads', methods=['POST'])
def cleanup_uploads():
    folder = os.path.join(Config.BASE_DIR, 'static', 'uploads')
    if os.path.exists(folder):
        shutil.rmtree(folder)
        os.makedirs(folder)
    return jsonify({'message': '✅ 已删除所有本地的项目上传文件！'})

# todo:新增删除项目总表
@app.route('/admin/cleanup_downloads', methods=['POST'])
def cleanup_downloads():
    stats_folder = os.path.join(Config.BASE_DIR, 'static', 'downloads', 'Statistics')
    regs_folder = os.path.join(Config.BASE_DIR, 'static', 'downloads', 'Registrations')
    project_summary_folder = os.path.join(Config.BASE_DIR, 'static', 'downloads', 'ProjectsST')  # ✅ 新增路径

    for folder in [stats_folder, regs_folder, project_summary_folder]:
        if os.path.exists(folder):
            for filename in os.listdir(folder):
                file_path = os.path.join(folder, filename)
                try:
                    if os.path.isfile(file_path):
                        os.remove(file_path)
                except Exception as e:
                    print(f"⚠️ 无法删除文件 {file_path}: {e}")

    return jsonify({'message': '✅ 已清空登记表、统计表和项目总表内容，但保留目录结构'})


# @app.route('/clear_form_session', methods=['POST'])
# def clear_form_session():
#     session.pop('project_form_data', None)
#     return '', 204  # 返回 No Content

# import subprocess
#
# def try_restart_supervisor():
#     try:
#         result = subprocess.run(
#             ['/usr/bin/sudo', '/home/ecs-user/restart_frontcourt.sh'],  # 调用你创建的脚本
#             check=True,
#             stdout=subprocess.PIPE,
#             stderr=subprocess.PIPE,
#             universal_newlines=True  # 等价于 text=True，兼容旧版本
#         )
#         return True, "✅ 应用已重启：" + result.stdout.strip()
#     except subprocess.CalledProcessError as e:
#         return False, f"❌ 重启失败，stderr: {e.stderr.strip() or '（无输出）'}"
#     except Exception as ex:
#         return False, f"❌ 出现未知错误：{str(ex)}"


# 恢复数据库从oss到本地操作确认机制
# @app.route('/admin/restore_db', methods=['POST'])
# def restore_db():
#     from datetime import datetime
#     import platform
#
#     now = datetime.now()
#     if not (0 <= now.hour < 15):
#         return jsonify({'status': 'fail', 'msg': '❌ 只能在凌晨0:00–15:00之间恢复数据库，以防误操作！'})
#
#     try:
#         from oss import restore_db_from_oss
#         restore_db_from_oss()
#
#         if platform.system().lower() != 'windows':
#             # ✅ 尝试重启
#             success, msg = try_restart_supervisor()
#             return jsonify({'status': 'success' if success else 'fail', 'msg': msg})
#         else:
#             return jsonify({'status': 'success', 'msg': '✅ 数据库恢复成功（开发环境无需重启）'})
#
#     except Exception as e:
#         return jsonify({'status': 'fail', 'msg': f'❌ 操作失败：{str(e)}'})

# 要不然就这样，逻辑如下：既然发送文件会一定要用到uploads包里的文件，那干脆，


@app.route('/edit_sub_project/<int:sub_project_id>', methods=['GET', 'POST'])
def edit_sub_project(sub_project_id):
    if not session.get('admin'):
        return redirect(url_for('login'))

    sub_project = SubProject.query.get_or_404(sub_project_id)

    if request.method == 'POST':
        # 修改字段
        sub_project.segment_name = request.form['segment_name']
        sub_project.start_time = datetime.strptime(request.form['start_time'], '%Y-%m-%dT%H:%M')
        sub_project.deadline = datetime.strptime(request.form['deadline'], '%Y-%m-%dT%H:%M')
        sub_project.deposit_amount = float(request.form['deposit_amount'])
        db.session.commit()
        flash("✅ 标段已更新", "success")
        return redirect(url_for('admin_panel'))

    return render_template("edit_sub_project.html", sub_project=sub_project)



# todo : 导出总项目表
@app.route('/export_project_summary')
def export_project_summary():
    if not session.get('admin'):
        return redirect(url_for('login'))

    projects = Project.query.order_by(Project.id.asc()).all()
    if not projects:
        flash("⚠️ 当前没有任何项目，无法导出", "warning")
        return redirect(url_for('admin_panel'))

    # ✅ 创建 Excel 工作簿
    wb = Workbook()
    ws = wb.active
    ws.title = "项目统计表"

    # ✅ 设置列宽
    ws.column_dimensions['A'].width = 8     # 序号
    ws.column_dimensions['B'].width = 50    # 项目名称
    ws.column_dimensions['C'].width = 20    # 项目编号
    ws.column_dimensions['D'].width = 20    # 采购人
    ws.column_dimensions['E'].width = 20    # 采购金额

    # ✅ 合并表头并格式化
    ws.merge_cells('A1:E1')
    ws['A1'] = "项目统计表"
    ws['A1'].font = Font(size=14, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')

    # ✅ 第二行标题
    headers = ["序号", "项目名称", "项目编号", "采购人", "采购金额"]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col_num, value=header)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center')

    # ✅ 填充数据内容
    for idx, project in enumerate(projects, 1):
        ws.cell(row=2 + idx, column=1, value=idx).alignment = Alignment(horizontal='center')
        ws.cell(row=2 + idx, column=2, value=project.name).alignment = Alignment(horizontal='center')
        ws.cell(row=2 + idx, column=3, value=project.code).alignment = Alignment(horizontal='center')
        ws.cell(row=2 + idx, column=4, value=project.purchaser or "").alignment = Alignment(horizontal='center')
        ws.cell(row=2 + idx, column=5, value=project.budget_amount or 0).alignment = Alignment(horizontal='center')

    # ✅ 添加筛选器
    ws.auto_filter = AutoFilter(ref=f"A2:E{2 + len(projects)}")

    # ✅ 文件名与路径
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    file_name = f"项目统计表_{timestamp}.xlsx"
    local_dir = os.path.join(Config.BASE_DIR, "static", "downloads", "ProjectsST")
    os.makedirs(local_dir, exist_ok=True)
    file_path = os.path.join(local_dir, file_name)

    # ✅ 保存本地 & 上传 OSS
    wb.save(file_path)
    upload_file_to_oss(file_path, f"项目总表/{file_name}")

    # ✅ 返回文件给浏览器下载
    return sf(file_path, as_attachment=True)

# 定时启动
'''
1. from flask_apscheduler import APScheduler：引入APScheduler。

2. class ConfigWithScheduler(Config):：创建一个带 SCHEDULER_API_ENABLED 的子配置类（为了让Flask-APScheduler能正常运行）。

3. scheduler = APScheduler()：实例化调度器。

4. app.config.from_object(ConfigWithScheduler)：让Flask加载这个新配置（一定要换掉原来的Config，否则调度器初始化会报错！）

5. with app.app_context(): init_db()：建表。

6. scheduler.init_app(app)：初始化调度器。

7. scheduler.start()：启动调度器。

8. app.run(debug=True, port=8000)：最后启动你的Flask服务。
'''


class ConfigWithScheduler(Config):
    SCHEDULER_API_ENABLED = True


scheduler = APScheduler()

# ================== 启动入口 ==================
if __name__ == "__main__":
    # ✅ 加上这段用于创建表（只执行一次即可）
    # with app.app_context():
    #     db.create_all()
    #     print("✅ MySQL 表结构已创建完毕")
    # with app.app_context():
    #     db.create_all()
    #
    #     # 检查是否已有管理员
    #     if not Admin.query.first():
    #         default_admin = Admin(
    #             username=app.config.get("ADMIN_USERNAME", "yngh"),
    #             password_hash=app.config.get("ADMIN_PASSWORD", generate_password_hash("yngh123")),
    #             email=app.config.get("ADMIN_EMAIL", "shaun7565@163.com")
    #         )
    #         db.session.add(default_admin)
    #         db.session.commit()
    #         print("✅ 已创建默认管理员账户")
    #     else:
    #         print("ℹ️ 管理员账户已存在，无需创建")

    app.config.from_object(ConfigWithScheduler)  # ✅ 注意！要换成带Scheduler配置的Config
    app.secret_key = Config.SECRET_KEY

    with app.app_context():
        init_db()
    scheduler.init_app(app)

    # 💡 5. 注册定时任务
    if os.environ.get('WERKZEUG_RUN_MAIN') == 'true':  # ✅ 只在主进程注册定时任务！
        # 定时发邮件
        scheduler.add_job(
            id='send_daily_project_summary',
            func=send_daily_project_summary,
            trigger='cron',
            hour=17,
            minute=0,
            timezone='Asia/Shanghai'
        )

        # 添加定时上传数据库任务（每天凌晨2点）
        scheduler.add_job(
            id='upload_db_to_oss',
            func=upload_db_to_oss,
            trigger='cron',
            hour=2,
            minute=0,
            timezone='Asia/Shanghai'
        )

        # 添加定时同步 static 下 OSS 文件任务（每天凌晨2:10）
        scheduler.add_job(
            id='sync_static_to_oss',
            func=sync_static_to_oss,
            trigger='cron',
            hour=2,
            minute=10,
            timezone='Asia/Shanghai'
        )

        scheduler.start()

    app.run(debug=True, port=8000, host="0.0.0.0")
